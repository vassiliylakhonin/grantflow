# grantflow/exporters/word_builder.py

from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from grantflow.exporters.donor_contracts import evaluate_export_contract
from grantflow.exporters.template_profile import (
    build_export_template_profile,
    resolve_export_template_key,
)
from grantflow.exporters.toc_normalization import normalize_toc_for_export, unwrap_toc_payload


def _compact_export_text(value: Any, *, max_len: int | None = None) -> str:
    text = " ".join(str(value or "").split()).strip()
    if not text:
        return ""
    for token in (
        "Evidence hint:",
        "Grounding gate warning:",
        "architect_retrieval_no_hits",
    ):
        idx = text.lower().find(token.lower())
        if idx >= 0:
            text = text[:idx].rstrip(" .,-")
    text = text.replace(" service delivery delivery", " service delivery")
    text = text.replace("  ", " ").strip()
    if max_len and len(text) > max_len:
        return f"{text[: max_len - 3].rstrip()}..."
    return text


def _normalized_indicator_rows(logframe_draft: Optional[Dict[str, Any]]) -> list[Dict[str, Any]]:
    if not isinstance(logframe_draft, dict):
        return []
    raw = logframe_draft.get("indicators")
    return [row for row in raw if isinstance(row, dict)] if isinstance(raw, list) else []


def _indicator_focus_rows(
    indicators: list[Dict[str, Any]],
    *,
    result_level: Optional[str] = None,
    limit: int = 2,
) -> list[Dict[str, Any]]:
    rows = indicators
    if result_level:
        token = str(result_level).strip().lower()
        rows = [row for row in indicators if str(row.get("result_level") or "").strip().lower() == token]
    return rows[: max(0, limit)]


def _add_indicator_focus_block(
    doc: Document,
    *,
    indicators: list[Dict[str, Any]],
    label: str = "Monitoring focus",
) -> None:
    focus_rows = [row for row in indicators if isinstance(row, dict)]
    if not focus_rows:
        return
    doc.add_paragraph(f"{label}:", style="List Bullet")
    for row in focus_rows:
        name = str(row.get("name") or "").strip() or str(row.get("indicator_id") or "Indicator")
        mov = str(row.get("means_of_verification") or "").strip()
        owner = str(row.get("owner") or "").strip()
        line = name
        if mov:
            line += f" | Means of verification: {mov}"
        if owner:
            line += f" | Owner: {owner}"
        doc.add_paragraph(line, style="List Bullet")


def _add_indicator_logframe_block(
    doc: Document,
    *,
    indicators: list[Dict[str, Any]],
    label: str = "Suggested LogFrame rows",
) -> None:
    focus_rows = [row for row in indicators if isinstance(row, dict)]
    if not focus_rows:
        return
    doc.add_paragraph(f"{label}:", style="List Bullet")
    for row in focus_rows:
        name = str(row.get("name") or "").strip() or str(row.get("indicator_id") or "Indicator")
        baseline = str(row.get("baseline") or "").strip() or "TBD"
        target = str(row.get("target") or "").strip() or "TBD"
        frequency = str(row.get("frequency") or "").strip()
        formula = str(row.get("formula") or "").strip()
        definition = " ".join(str(row.get("definition") or "").split()).strip()
        justification = " ".join(str(row.get("justification") or "").split()).strip()
        line = f"{name} | Baseline/Target: {baseline} -> {target}"
        if frequency:
            line += f" | Frequency: {frequency}"
        if formula:
            line += f" | Formula: {formula}"
        if definition:
            line += f" | Result focus: {definition[:110].rstrip()}{'...' if len(definition) > 110 else ''}"
        if justification:
            line += f" | Measurement intent: {justification[:110].rstrip()}{'...' if len(justification) > 110 else ''}"
        doc.add_paragraph(line, style="List Bullet")


def _review_readiness_rows(
    *,
    quality_summary: dict[str, Any],
    citations: list[Dict[str, Any]],
    critic_findings: list[Dict[str, Any]],
    review_comments: list[Dict[str, Any]],
) -> list[tuple[str, Any]]:
    open_findings = [
        item
        for item in critic_findings
        if isinstance(item, dict) and str(item.get("status") or "open").strip().lower() not in {"resolved", "closed"}
    ]
    finding_ack_queue_count = sum(
        1
        for item in critic_findings
        if isinstance(item, dict) and str(item.get("status") or "open").strip().lower() == "open"
    )
    finding_resolve_queue_count = sum(
        1
        for item in critic_findings
        if isinstance(item, dict) and str(item.get("status") or "").strip().lower() == "acknowledged"
    )
    high_findings = [
        item
        for item in open_findings
        if str(item.get("severity") or "").strip().lower() in {"high", "critical", "fatal"}
    ]
    open_comments = [
        item
        for item in review_comments
        if isinstance(item, dict) and str(item.get("status") or "open").strip().lower() not in {"resolved", "closed"}
    ]
    low_confidence = [
        item
        for item in citations
        if isinstance(item, dict) and str(item.get("citation_type") or "").strip().lower() == "rag_low_confidence"
    ]
    fallback = [
        item
        for item in citations
        if isinstance(item, dict)
        and str(item.get("citation_type") or "").strip().lower() in {"fallback_namespace", "strategy_reference"}
    ]
    rows = [
        ("Needs revision", quality_summary.get("needs_revision")),
        ("Open critic findings", len(open_findings)),
        (
            "Acknowledged critic findings",
            len(
                [
                    item
                    for item in critic_findings
                    if isinstance(item, dict) and str(item.get("status") or "").strip().lower() == "acknowledged"
                ]
            ),
        ),
        (
            "Resolved critic findings",
            len(
                [
                    item
                    for item in critic_findings
                    if isinstance(item, dict)
                    and str(item.get("status") or "").strip().lower() in {"resolved", "closed"}
                ]
            ),
        ),
        ("High-severity open findings", len(high_findings)),
        ("Open review comments", len(open_comments)),
        ("Low-confidence citations", len(low_confidence)),
        ("Fallback/strategy citations", len(fallback)),
    ]
    top_actions: list[str] = []
    for item in open_findings:
        if not isinstance(item, dict):
            continue
        action = str(item.get("reviewer_next_step") or item.get("recommended_action") or "").strip()
        if action and action not in top_actions:
            top_actions.append(action)
        if len(top_actions) >= 2:
            break
    for idx, action in enumerate(top_actions, start=1):
        rows.append((f"Top reviewer action {idx}", action))
    acknowledged_comments = [
        item
        for item in review_comments
        if isinstance(item, dict) and str(item.get("status") or "").strip().lower() == "acknowledged"
    ]
    resolved_comments = [
        item
        for item in review_comments
        if isinstance(item, dict) and str(item.get("status") or "").strip().lower() in {"resolved", "closed"}
    ]
    total_items = len(critic_findings) + len(review_comments)
    resolved_items = len(
        [
            item
            for item in critic_findings
            if isinstance(item, dict) and str(item.get("status") or "").strip().lower() in {"resolved", "closed"}
        ]
    ) + len(resolved_comments)
    acknowledged_items = len(
        [
            item
            for item in critic_findings
            if isinstance(item, dict) and str(item.get("status") or "").strip().lower() == "acknowledged"
        ]
    ) + len(acknowledged_comments)
    rows.extend(
        [
            ("Acknowledged review comments", len(acknowledged_comments)),
            ("Resolved review comments", len(resolved_comments)),
            (
                "Critic finding resolution rate",
                (
                    round(
                        len(
                            [
                                item
                                for item in critic_findings
                                if isinstance(item, dict)
                                and str(item.get("status") or "").strip().lower() in {"resolved", "closed"}
                            ]
                        )
                        / len(critic_findings),
                        4,
                    )
                    if critic_findings
                    else None
                ),
            ),
            (
                "Critic finding acknowledgment rate",
                (
                    round(
                        (
                            len(
                                [
                                    item
                                    for item in critic_findings
                                    if isinstance(item, dict)
                                    and str(item.get("status") or "").strip().lower() in {"resolved", "closed"}
                                ]
                            )
                            + len(
                                [
                                    item
                                    for item in critic_findings
                                    if isinstance(item, dict)
                                    and str(item.get("status") or "").strip().lower() == "acknowledged"
                                ]
                            )
                        )
                        / len(critic_findings),
                        4,
                    )
                    if critic_findings
                    else None
                ),
            ),
            ("Reviewer workflow resolution rate", round(resolved_items / total_items, 4) if total_items else None),
            (
                "Reviewer workflow acknowledgment rate",
                round((resolved_items + acknowledged_items) / total_items, 4) if total_items else None,
            ),
            (
                "Next primary review action",
                (
                    "ack_finding"
                    if finding_ack_queue_count > 0
                    else (
                        "resolve_finding"
                        if finding_resolve_queue_count > 0
                        else (
                            "ack_comment"
                            if len(open_comments) > 0
                            else (
                                "resolve_comment"
                                if len(acknowledged_comments) > 0
                                else ("reopen_comment" if len(resolved_comments) > 0 else None)
                            )
                        )
                    )
                ),
            ),
            ("Finding ack queue", finding_ack_queue_count),
            ("Finding resolve queue", finding_resolve_queue_count),
            ("Comment ack queue", len(open_comments)),
            ("Comment resolve queue", len(acknowledged_comments)),
            ("Comment reopen queue", len(resolved_comments)),
        ]
    )

    now_candidates = []
    for item in review_comments:
        if not isinstance(item, dict):
            continue
        for key in ("last_transition_at", "updated_ts", "resolved_at", "acknowledged_at", "ts", "due_at"):
            value = item.get(key)
            if value:
                now_candidates.append(value)
    # lightweight aging summary for exports without importing API helpers
    from datetime import datetime, timezone

    def _parse_dt(value: Any):
        if not value:
            return None
        try:
            text = str(value).replace("Z", "+00:00")
            dt = datetime.fromisoformat(text)
            if dt.tzinfo is None:
                return dt.replace(tzinfo=timezone.utc)
            return dt.astimezone(timezone.utc)
        except Exception:
            return None

    ref = max((_parse_dt(value) for value in now_candidates), default=None)
    if ref is None:
        ref = datetime.now(timezone.utc)
    aging = {"3-7d": 0, ">7d": 0}
    for item in review_comments:
        if not isinstance(item, dict):
            continue
        status = str(item.get("status") or "open").strip().lower()
        if status in {"resolved", "closed"}:
            continue
        transition = (
            _parse_dt(item.get("last_transition_at")) or _parse_dt(item.get("updated_ts")) or _parse_dt(item.get("ts"))
        )
        if transition is None:
            continue
        hours = max(0.0, (ref - transition).total_seconds() / 3600.0)
        if 72.0 <= hours < 168.0:
            aging["3-7d"] += 1
        elif hours >= 168.0:
            aging[">7d"] += 1
    rows.extend(
        [
            ("Comment threads aged 3-7d", aging["3-7d"]),
            ("Comment threads aged >7d", aging[">7d"]),
        ]
    )
    return [(label, value) for label, value in rows if value is not None and value != ""]


def _citation_summary_line(citation: Dict[str, Any]) -> str:
    stage = citation.get("stage", "")
    ctype = citation.get("citation_type", "")
    used_for = citation.get("used_for", "")
    label = citation.get("label") or citation.get("source") or citation.get("namespace") or ""
    page = citation.get("page")
    chunk = citation.get("chunk")
    confidence = citation.get("citation_confidence")
    statement_path = citation.get("statement_path")
    result_level = citation.get("result_level")
    evidence_signal = citation.get("evidence_signal")

    parts = []
    if stage:
        parts.append(f"[{stage}]")
    if ctype:
        parts.append(str(ctype))
    if used_for:
        parts.append(f"for {used_for}")
    if label:
        parts.append(f"- {label}")
    if confidence is not None:
        try:
            parts.append(f"(conf {float(confidence):.2f})")
        except (TypeError, ValueError):
            parts.append(f"(conf {confidence})")
    if result_level:
        parts.append(f"(level {result_level})")
    if statement_path:
        parts.append(f"(path {statement_path})")
    if evidence_signal:
        parts.append(f"(evidence {evidence_signal})")
    if page is not None:
        parts.append(f"(p.{page})")
    if chunk is not None:
        parts.append(f"(chunk {chunk})")
    return " ".join(parts).strip()


def _add_citation_traceability_section(doc: Document, citations: list[Dict[str, Any]]) -> None:
    if not citations:
        return

    doc.add_heading("Citation Traceability", level=1)
    doc.add_paragraph(
        "Structured citation records collected during drafting (strategy references and/or RAG retrieval results)."
    )
    for citation in citations:
        doc.add_paragraph(_citation_summary_line(citation), style="List Bullet")
        review_hint = str(citation.get("review_hint") or "").strip()
        if review_hint:
            p = doc.add_paragraph()
            p.add_run(f"Review hint: {review_hint[:240]}").italic = True
        excerpt = (citation.get("excerpt") or "").strip()
        if excerpt:
            p = doc.add_paragraph()
            p.add_run(f"Excerpt: {excerpt[:240]}").italic = True


def _add_critic_findings_section(doc: Document, critic_findings: list[Dict[str, Any]]) -> None:
    if not critic_findings:
        return

    doc.add_heading("Critic Findings", level=1)
    doc.add_paragraph("Rule-based and/or LLM critic findings captured during review.")
    for finding in critic_findings:
        status = finding.get("status") or "open"
        severity = finding.get("severity") or "unknown"
        code = finding.get("code") or "FINDING"
        section = finding.get("section") or "general"
        review_title = str(finding.get("review_title") or "").strip()
        title_suffix = review_title or str(code)
        title = f"[{status}] [{severity}] {section} · {title_suffix}"
        doc.add_paragraph(title, style="List Bullet")
        message = str(finding.get("message") or "").strip()
        if message:
            doc.add_paragraph(message)
        recommended_action = str(finding.get("recommended_action") or "").strip()
        if recommended_action:
            p = doc.add_paragraph()
            p.add_run(f"Recommended action: {recommended_action}").italic = True
        reviewer_next_step = str(finding.get("reviewer_next_step") or "").strip()
        if reviewer_next_step:
            p = doc.add_paragraph()
            p.add_run(f"Reviewer next step: {reviewer_next_step}").italic = True
        fix_hint = str(finding.get("fix_hint") or "").strip()
        if fix_hint:
            p = doc.add_paragraph()
            p.add_run(f"Fix hint: {fix_hint}").italic = True
        finding_id = finding.get("id") or finding.get("finding_id")
        meta_bits: list[str] = []
        if finding.get("review_bucket"):
            meta_bits.append(f"bucket={finding.get('review_bucket')}")
        if finding.get("triage_priority"):
            meta_bits.append(f"priority={finding.get('triage_priority')}")
        if finding.get("version_id"):
            meta_bits.append(f"version={finding.get('version_id')}")
        if finding_id:
            meta_bits.append(f"finding_id={finding_id}")
        if finding.get("source"):
            meta_bits.append(f"source={finding.get('source')}")
        if meta_bits:
            doc.add_paragraph(" · ".join(meta_bits))


def _add_review_comments_section(doc: Document, review_comments: list[Dict[str, Any]]) -> None:
    if not review_comments:
        return

    doc.add_heading("Review Comments", level=1)
    doc.add_paragraph("Reviewer notes and resolution workflow comments linked to draft versions/findings.")
    for comment in review_comments:
        status = comment.get("status") or "open"
        section = comment.get("section") or "general"
        author = comment.get("author") or "reviewer"
        title = f"[{status}] {section} · {author}"
        doc.add_paragraph(title, style="List Bullet")
        message = str(comment.get("message") or "").strip()
        if message:
            doc.add_paragraph(message)
        meta_bits: list[str] = []
        if comment.get("version_id"):
            meta_bits.append(f"version={comment.get('version_id')}")
        if comment.get("linked_finding_id"):
            meta_bits.append(f"linked_finding_id={comment.get('linked_finding_id')}")
        if comment.get("ts"):
            meta_bits.append(str(comment.get("ts")))
        if meta_bits:
            doc.add_paragraph(" · ".join(meta_bits))


def _add_quality_summary_section(doc: Document, quality_summary: dict[str, Any]) -> None:
    if not quality_summary:
        return

    rows = [
        ("Quality score", quality_summary.get("quality_score")),
        ("Critic score", quality_summary.get("critic_score")),
        ("Needs revision", quality_summary.get("needs_revision")),
        ("Critic engine", quality_summary.get("engine")),
        ("Rule score", quality_summary.get("rule_score")),
        ("LLM score", quality_summary.get("llm_score")),
        ("Fatal flaw count", quality_summary.get("fatal_flaw_count")),
        ("Citation count", quality_summary.get("citation_count")),
    ]
    present_rows = [(label, value) for label, value in rows if value is not None and value != ""]
    if not present_rows:
        return

    doc.add_heading("Quality Summary", level=1)
    doc.add_paragraph("Snapshot of current draft quality, critic status, and citation volume at export time.")
    for label, value in present_rows:
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")


def _add_review_readiness_section(
    doc: Document,
    *,
    quality_summary: dict[str, Any],
    citations: list[Dict[str, Any]],
    critic_findings: list[Dict[str, Any]],
    review_comments: list[Dict[str, Any]],
) -> None:
    rows = _review_readiness_rows(
        quality_summary=quality_summary,
        citations=citations,
        critic_findings=critic_findings,
        review_comments=review_comments,
    )
    if not rows:
        return

    doc.add_heading("Review Readiness", level=1)
    doc.add_paragraph("Compact reviewer snapshot for triage before sign-off or export handoff.")
    for label, value in rows:
        doc.add_paragraph(f"{label}: {value}", style="List Bullet")


def _toc_root(toc_draft: Dict[str, Any], donor_id: str | None = None) -> Dict[str, Any]:
    if not donor_id:
        return unwrap_toc_payload(toc_draft)
    donor_key = resolve_export_template_key(donor_id=donor_id, toc_payload=toc_draft)
    return normalize_toc_for_export(donor_key, unwrap_toc_payload(toc_draft))


def _render_usaid_toc(doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=2)
    impact_focus = _indicator_focus_rows(indicators, result_level="impact", limit=1)
    doc.add_heading("USAID Results Framework", level=1)
    goal = _compact_export_text(toc.get("project_goal"))
    if goal:
        doc.add_heading("Project Goal", level=2)
        doc.add_paragraph(goal)
        _add_indicator_focus_block(doc, indicators=impact_focus, label="Suggested learning focus")
        _add_indicator_logframe_block(doc, indicators=impact_focus, label="Suggested goal-level LogFrame rows")

    do_list = toc.get("development_objectives")
    if isinstance(do_list, list) and do_list:
        doc.add_heading("Development Objectives", level=2)
        for do in do_list:
            if not isinstance(do, dict):
                continue
            do_id = str(do.get("do_id") or "").strip()
            do_title = _compact_export_text(do.get("description"))
            doc.add_heading(f"{do_id or 'DO'} — {do_title or 'Development Objective'}", level=3)
            _add_indicator_focus_block(doc, indicators=outcome_focus, label="Suggested performance monitoring focus")
            _add_indicator_logframe_block(doc, indicators=outcome_focus, label="Suggested performance indicator rows")
            ir_list = do.get("intermediate_results")
            if not isinstance(ir_list, list):
                continue
            for ir in ir_list:
                if not isinstance(ir, dict):
                    continue
                ir_id = str(ir.get("ir_id") or "").strip()
                ir_desc = _compact_export_text(ir.get("description"))
                doc.add_paragraph(f"IR: {ir_id or '-'} — {ir_desc or '-'}", style="List Bullet")
                outputs = ir.get("outputs")
                if not isinstance(outputs, list):
                    continue
                for out in outputs:
                    if not isinstance(out, dict):
                        continue
                    output_id = str(out.get("output_id") or "").strip()
                    output_desc = _compact_export_text(out.get("description"))
                    doc.add_paragraph(f"Output: {output_id or '-'} — {output_desc or '-'}", style="List Bullet")
                    indicators_raw = out.get("indicators")
                    if not isinstance(indicators_raw, list):
                        continue
                    output_indicators: list[dict[str, Any]] = [
                        indicator for indicator in indicators_raw if isinstance(indicator, dict)
                    ]
                    for indicator in output_indicators:
                        code = str(indicator.get("indicator_code") or "").strip()
                        name = _compact_export_text(indicator.get("name"))
                        target = str(indicator.get("target") or "").strip()
                        citation = str(indicator.get("citation") or "").strip()
                        line = f"Indicator: {code or '-'} — {name or '-'}"
                        if target:
                            line += f" | Target: {target}"
                        if citation:
                            line += f" | Citation: {citation}"
                        doc.add_paragraph(line, style="List Bullet")

    assumptions = toc.get("critical_assumptions")
    if isinstance(assumptions, list) and assumptions:
        doc.add_heading("Critical Assumptions", level=2)
        for assumption in assumptions:
            doc.add_paragraph(str(assumption), style="List Bullet")


def _render_eu_toc(doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=2)
    output_focus = _indicator_focus_rows(indicators, result_level="output", limit=2)
    doc.add_heading("EU Intervention Logic", level=1)
    overall = toc.get("overall_objective")
    rendered = False
    if isinstance(overall, dict):
        objective_id = str(overall.get("objective_id") or "").strip()
        title = _compact_export_text(overall.get("title"))
        rationale = _compact_export_text(overall.get("rationale"))
        doc.add_heading("Overall Objective", level=2)
        if objective_id or title:
            doc.add_paragraph(f"{objective_id or 'Objective'} — {title or '-'}")
        if rationale:
            doc.add_paragraph(rationale)
        _add_indicator_focus_block(doc, indicators=outcome_focus, label="Suggested monitoring focus")
        _add_indicator_logframe_block(doc, indicators=outcome_focus, label="Suggested objective-level LogFrame rows")
        rendered = True

    specific_objectives = toc.get("specific_objectives")
    if isinstance(specific_objectives, list) and specific_objectives:
        doc.add_heading("Specific Objectives", level=2)
        for row in specific_objectives:
            if not isinstance(row, dict):
                continue
            objective_id = str(row.get("objective_id") or "").strip()
            title = _compact_export_text(row.get("title"))
            rationale = _compact_export_text(row.get("rationale"))
            doc.add_paragraph(f"{objective_id or 'SO'} — {title or '-'}", style="List Bullet")
            if rationale:
                doc.add_paragraph(rationale)
            _add_indicator_focus_block(doc, indicators=outcome_focus[:1], label="Suggested verification focus")
            _add_indicator_logframe_block(
                doc,
                indicators=outcome_focus[:1],
                label="Suggested specific-objective indicator rows",
            )
        rendered = True

    expected_outcomes = toc.get("expected_outcomes")
    if isinstance(expected_outcomes, list) and expected_outcomes:
        doc.add_heading("Expected Outcomes", level=2)
        for row in expected_outcomes:
            if not isinstance(row, dict):
                continue
            outcome_id = str(row.get("outcome_id") or "").strip()
            title = _compact_export_text(row.get("title"))
            expected_change = _compact_export_text(row.get("expected_change"))
            doc.add_paragraph(f"{outcome_id or 'Outcome'} — {title or '-'}", style="List Bullet")
            if expected_change:
                doc.add_paragraph(expected_change)
            _add_indicator_focus_block(
                doc, indicators=output_focus[:1] or outcome_focus[:1], label="Suggested delivery focus"
            )
            _add_indicator_logframe_block(
                doc,
                indicators=output_focus[:1] or outcome_focus[:1],
                label="Suggested outcome indicator rows",
            )
        rendered = True

    assumptions = toc.get("assumptions")
    if isinstance(assumptions, list) and assumptions:
        doc.add_heading("Assumptions", level=2)
        for assumption in assumptions:
            doc.add_paragraph(str(assumption), style="List Bullet")
        rendered = True

    risks = toc.get("risks")
    if isinstance(risks, list) and risks:
        doc.add_heading("Risks", level=2)
        for risk in risks:
            doc.add_paragraph(str(risk), style="List Bullet")
        rendered = True

    if not rendered:
        doc.add_paragraph("Overall objective is not provided in EU schema draft.")


def _render_worldbank_toc(
    doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None
) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=2)
    impact_focus = _indicator_focus_rows(indicators, result_level="impact", limit=1)
    doc.add_heading("World Bank Results Framework", level=1)
    rendered = False
    pdo = _compact_export_text(toc.get("project_development_objective"))
    if pdo:
        doc.add_heading("Project Development Objective (PDO)", level=2)
        doc.add_paragraph(pdo)
        _add_indicator_focus_block(
            doc, indicators=impact_focus or outcome_focus[:1], label="Suggested PDO monitoring focus"
        )
        _add_indicator_logframe_block(
            doc,
            indicators=impact_focus or outcome_focus[:1],
            label="Suggested PDO indicator rows",
        )
        rendered = True

    objectives = toc.get("objectives")
    if isinstance(objectives, list) and objectives:
        doc.add_heading("Objectives", level=2)
        for obj in objectives:
            if not isinstance(obj, dict):
                continue
            objective_id = str(obj.get("objective_id") or "").strip()
            title = _compact_export_text(obj.get("title"))
            description = _compact_export_text(obj.get("description"))
            doc.add_heading(f"{objective_id or 'Objective'} — {title or '-'}", level=3)
            _add_indicator_focus_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested results monitoring focus",
            )
            _add_indicator_logframe_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested objective indicator rows",
            )
            if description:
                doc.add_paragraph(description)
        rendered = True

    results_chain = toc.get("results_chain")
    if isinstance(results_chain, list) and results_chain:
        doc.add_heading("Results Chain", level=2)
        for row in results_chain:
            if not isinstance(row, dict):
                continue
            result_id = str(row.get("result_id") or "").strip()
            title = _compact_export_text(row.get("title"))
            description = _compact_export_text(row.get("description"))
            indicator_focus = _compact_export_text(row.get("indicator_focus"))
            doc.add_paragraph(f"{result_id or 'Result'} — {title or '-'}", style="List Bullet")
            if description:
                doc.add_paragraph(description)
            if indicator_focus:
                p = doc.add_paragraph()
                p.add_run(f"Indicator focus: {indicator_focus}").italic = True
            _add_indicator_focus_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested verification focus",
            )
            _add_indicator_logframe_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested result indicator rows",
            )
        rendered = True

    assumptions = toc.get("assumptions")
    if isinstance(assumptions, list) and assumptions:
        doc.add_heading("Assumptions", level=2)
        for assumption in assumptions:
            doc.add_paragraph(str(assumption), style="List Bullet")
        rendered = True

    risks = toc.get("risks")
    if isinstance(risks, list) and risks:
        doc.add_heading("Risks", level=2)
        for risk in risks:
            doc.add_paragraph(str(risk), style="List Bullet")
        rendered = True

    if not rendered:
        doc.add_paragraph("No World Bank objectives found in draft.")


def _render_giz_toc(doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=2)
    output_focus = _indicator_focus_rows(indicators, result_level="output", limit=2)
    doc.add_heading("GIZ Results & Sustainability Logic", level=1)
    rendered = False

    programme_objective = str(toc.get("programme_objective") or "").strip()
    if programme_objective:
        doc.add_heading("Programme Objective", level=2)
        doc.add_paragraph(programme_objective)
        _add_indicator_focus_block(
            doc,
            indicators=output_focus[:1] or outcome_focus[:1],
            label="Suggested implementation monitoring focus",
        )
        _add_indicator_logframe_block(
            doc,
            indicators=output_focus[:1] or outcome_focus[:1],
            label="Suggested programme-objective indicator rows",
        )
        rendered = True

    outputs = toc.get("outputs")
    if isinstance(outputs, list) and outputs:
        doc.add_heading("Outputs", level=2)
        for output in outputs:
            doc.add_paragraph(str(output), style="List Bullet")
        _add_indicator_focus_block(
            doc,
            indicators=output_focus[:1] or outcome_focus[:1],
            label="Suggested delivery verification focus",
        )
        _add_indicator_logframe_block(
            doc,
            indicators=output_focus[:1] or outcome_focus[:1],
            label="Suggested output indicator rows",
        )
        rendered = True

    outcomes = toc.get("outcomes")
    if isinstance(outcomes, list) and outcomes:
        doc.add_heading("Outcomes", level=2)
        for outcome in outcomes:
            if not isinstance(outcome, dict):
                continue
            title = str(outcome.get("title") or "").strip()
            description = str(outcome.get("description") or "").strip()
            partner_role = str(outcome.get("partner_role") or "").strip()
            doc.add_paragraph(title or "Outcome", style="List Bullet")
            if description:
                doc.add_paragraph(description)
            if partner_role:
                p = doc.add_paragraph()
                p.add_run(f"Partner role: {partner_role}").italic = True
            _add_indicator_focus_block(
                doc,
                indicators=output_focus[:1] or outcome_focus[:1],
                label="Suggested sustainability monitoring focus",
            )
            _add_indicator_logframe_block(
                doc,
                indicators=output_focus[:1] or outcome_focus[:1],
                label="Suggested outcome indicator rows",
            )
        rendered = True

    sustainability_factors = toc.get("sustainability_factors")
    if isinstance(sustainability_factors, list) and sustainability_factors:
        doc.add_heading("Sustainability Factors", level=2)
        for factor in sustainability_factors:
            doc.add_paragraph(str(factor), style="List Bullet")
        rendered = True

    assumptions_risks = toc.get("assumptions_risks")
    if isinstance(assumptions_risks, list) and assumptions_risks:
        doc.add_heading("Assumptions & Risks", level=2)
        for item in assumptions_risks:
            doc.add_paragraph(str(item), style="List Bullet")
        rendered = True

    if not rendered:
        doc.add_paragraph("No GIZ-specific ToC structure found in draft.")


def _render_un_agencies_toc(
    doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None
) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=2)
    impact_focus = _indicator_focus_rows(indicators, result_level="impact", limit=1)
    doc.add_heading("UN Agency Program Logic", level=1)

    brief = str(toc.get("brief") or "").strip()
    if brief:
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(brief)
        _add_indicator_focus_block(
            doc,
            indicators=impact_focus or outcome_focus[:1],
            label="Suggested results focus",
        )
        _add_indicator_logframe_block(
            doc,
            indicators=impact_focus or outcome_focus[:1],
            label="Suggested overview indicator rows",
        )

    objectives = toc.get("objectives")
    if isinstance(objectives, list) and objectives:
        doc.add_heading("Development Objectives", level=2)
        for row in objectives:
            if not isinstance(row, dict):
                continue
            title = str(row.get("title") or "Untitled").strip()
            description = str(row.get("description") or "").strip()
            doc.add_heading(title, level=3)
            if description:
                doc.add_paragraph(description)
            _add_indicator_focus_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested monitoring focus",
            )
            _add_indicator_logframe_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested objective indicator rows",
            )
            citation = str(row.get("citation") or "").strip()
            if citation:
                p = doc.add_paragraph()
                p.add_run(f"Citation: {citation}").italic = True

    outcomes = toc.get("outcomes")
    if isinstance(outcomes, list) and outcomes:
        doc.add_heading("Expected Outcomes", level=2)
        for row in outcomes:
            if not isinstance(row, dict):
                continue
            title = str(row.get("title") or row.get("name") or "Outcome").strip()
            description = str(row.get("description") or row.get("expected_change") or "").strip()
            doc.add_paragraph(title, style="List Bullet")
            if description:
                doc.add_paragraph(description)
            _add_indicator_logframe_block(
                doc,
                indicators=outcome_focus[:1],
                label="Suggested outcome indicator rows",
            )

    toc_indicators = toc.get("indicators")
    if isinstance(toc_indicators, list) and toc_indicators:
        doc.add_heading("Key Indicators", level=2)
        for ind in toc_indicators:
            if not isinstance(ind, dict):
                continue
            doc.add_paragraph(f"• {ind.get('name', 'Unknown')}", style="List Bullet")
            justification = str(ind.get("justification") or "").strip()
            if justification:
                doc.add_paragraph(f"  Justification: {justification}", style="Intense Quote")


def _render_state_department_toc(
    doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None
) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=2)
    impact_focus = _indicator_focus_rows(indicators, result_level="impact", limit=1)
    doc.add_heading("U.S. Department of State Program Logic", level=1)
    rendered = False

    strategic_context = str(toc.get("strategic_context") or "").strip()
    if strategic_context:
        doc.add_heading("Strategic Context", level=2)
        doc.add_paragraph(strategic_context)
        rendered = True

    program_goal = str(toc.get("program_goal") or "").strip()
    if program_goal:
        doc.add_heading("Program Goal", level=2)
        doc.add_paragraph(program_goal)
        _add_indicator_focus_block(
            doc,
            indicators=impact_focus or outcome_focus[:1],
            label="Suggested strategic monitoring focus",
        )
        rendered = True

    objectives = toc.get("objectives")
    if isinstance(objectives, list) and objectives:
        doc.add_heading("Objectives", level=2)
        for row in objectives:
            if not isinstance(row, dict):
                continue
            objective = str(row.get("objective") or "").strip()
            line_of_effort = str(row.get("line_of_effort") or "").strip()
            expected_change = str(row.get("expected_change") or "").strip()
            title = objective or "Objective"
            if line_of_effort:
                title += f" ({line_of_effort})"
            doc.add_paragraph(title, style="List Bullet")
            _add_indicator_focus_block(
                doc,
                indicators=outcome_focus[:1] or impact_focus,
                label="Suggested delivery monitoring focus",
            )
            if expected_change:
                doc.add_paragraph(expected_change)
        rendered = True

    stakeholder_map = toc.get("stakeholder_map")
    if isinstance(stakeholder_map, list) and stakeholder_map:
        doc.add_heading("Stakeholder Map", level=2)
        for stakeholder in stakeholder_map:
            doc.add_paragraph(str(stakeholder), style="List Bullet")
        rendered = True

    risk_mitigation = toc.get("risk_mitigation")
    if isinstance(risk_mitigation, list) and risk_mitigation:
        doc.add_heading("Risk Mitigation", level=2)
        for risk in risk_mitigation:
            doc.add_paragraph(str(risk), style="List Bullet")
        rendered = True

    if not rendered:
        doc.add_paragraph("No State Department-specific ToC structure found in draft.")


def _render_generic_toc(doc: Document, toc: Dict[str, Any]) -> None:
    if "brief" in toc:
        doc.add_heading("Overview", level=1)
        doc.add_paragraph(str(toc.get("brief") or ""))

    if "objectives" in toc and isinstance(toc.get("objectives"), list):
        doc.add_heading("Development Objectives", level=1)
        for obj in toc["objectives"]:
            if not isinstance(obj, dict):
                continue
            doc.add_heading(str(obj.get("title") or "Untitled"), level=2)
            doc.add_paragraph(str(obj.get("description") or ""))
            citation = str(obj.get("citation") or "").strip()
            if citation:
                p = doc.add_paragraph()
                p.add_run(f"Citation: {citation}").italic = True

    if "indicators" in toc and isinstance(toc.get("indicators"), list):
        doc.add_heading("Key Indicators", level=1)
        for ind in toc["indicators"]:
            if not isinstance(ind, dict):
                continue
            doc.add_paragraph(f"• {ind.get('name', 'Unknown')}", style="List Bullet")
            if "justification" in ind:
                doc.add_paragraph(f"  Justification: {ind['justification']}", style="Intense Quote")


def _render_evaluation_rfq_toc(
    doc: Document, toc: Dict[str, Any], *, logframe_draft: Optional[Dict[str, Any]] = None
) -> None:
    indicators = _normalized_indicator_rows(logframe_draft)
    output_focus = _indicator_focus_rows(indicators, result_level="output", limit=2)
    outcome_focus = _indicator_focus_rows(indicators, result_level="outcome", limit=1)
    doc.add_heading("Evaluation RFQ Technical Proposal", level=1)

    brief = str(toc.get("brief") or "").strip()
    if brief:
        doc.add_heading("Assignment Summary", level=2)
        doc.add_paragraph(brief)

    background = str(toc.get("background_context") or "").strip()
    if background:
        doc.add_heading("Assignment Background", level=2)
        doc.add_paragraph(background)

    purpose = str(toc.get("evaluation_purpose") or "").strip()
    if purpose:
        doc.add_heading("Evaluation Purpose", level=2)
        doc.add_paragraph(purpose)

    questions = toc.get("evaluation_questions")
    if isinstance(questions, list) and questions:
        doc.add_heading("Evaluation Questions", level=2)
        for question in questions:
            doc.add_paragraph(str(question), style="List Bullet")

    methodology = str(toc.get("methodology_overview") or "").strip()
    if methodology:
        doc.add_heading("Methodology", level=2)
        doc.add_paragraph(methodology)
        _add_indicator_focus_block(
            doc,
            indicators=output_focus or outcome_focus,
            label="Suggested delivery monitoring focus",
        )

    components = toc.get("methodology_components")
    if isinstance(components, list) and components:
        doc.add_heading("Method Components", level=2)
        for row in components:
            if not isinstance(row, dict):
                continue
            title = str(row.get("method") or "Method").strip()
            purpose_text = str(row.get("purpose") or "").strip()
            respondent_group = str(row.get("respondent_group") or "").strip()
            evidence_source = str(row.get("evidence_source") or "").strip()
            doc.add_paragraph(title, style="List Bullet")
            if purpose_text:
                doc.add_paragraph(purpose_text)
            if respondent_group:
                doc.add_paragraph(f"Respondent group: {respondent_group}")
            if evidence_source:
                doc.add_paragraph(f"Evidence source: {evidence_source}")

    organization_information = str(toc.get("organization_information") or "").strip()
    if organization_information:
        doc.add_heading("Organization Information", level=2)
        doc.add_paragraph(organization_information)

    technical_approach = str(toc.get("technical_approach_summary") or "").strip()
    if technical_approach:
        doc.add_heading("Analysis and Proposed Approaches / Methodologies", level=2)
        doc.add_paragraph(technical_approach)

    sampling_plan = str(toc.get("sampling_plan") or "").strip()
    if sampling_plan:
        doc.add_heading("Sampling Plan", level=2)
        doc.add_paragraph(sampling_plan)

    software = toc.get("analytical_software")
    if isinstance(software, list) and software:
        doc.add_heading("Analytical Approach and Software", level=2)
        for item in software:
            doc.add_paragraph(str(item), style="List Bullet")

    ethics = toc.get("ethical_considerations")
    if isinstance(ethics, list) and ethics:
        doc.add_heading("Ethical Considerations", level=2)
        for item in ethics:
            doc.add_paragraph(str(item), style="List Bullet")

    team = toc.get("team_composition")
    if isinstance(team, list) and team:
        doc.add_heading("Personnel and Team Composition", level=2)
        for row in team:
            if not isinstance(row, dict):
                continue
            role = str(row.get("role") or "Role").strip()
            responsibility = str(row.get("responsibility") or "").strip()
            doc.add_paragraph(role, style="List Bullet")
            if responsibility:
                doc.add_paragraph(responsibility)

    key_personnel = toc.get("key_personnel")
    if isinstance(key_personnel, list) and key_personnel:
        doc.add_heading("Key Personnel and CV Readiness", level=2)
        for row in key_personnel:
            if not isinstance(row, dict):
                continue
            name = str(row.get("name") or "Personnel").strip()
            role = str(row.get("role") or "").strip()
            qualifications = str(row.get("qualifications") or "").strip()
            level_of_effort = str(row.get("level_of_effort") or "").strip()
            cv_status = str(row.get("cv_status") or "").strip()
            title = name if not role else f"{name} ({role})"
            doc.add_paragraph(title, style="List Bullet")
            details = []
            if qualifications:
                details.append(f"Qualifications: {qualifications}")
            if level_of_effort:
                details.append(f"Level of effort: {level_of_effort}")
            if cv_status:
                details.append(f"CV status: {cv_status}")
            if details:
                doc.add_paragraph(" | ".join(details))

    loe = str(toc.get("level_of_effort_summary") or "").strip()
    if loe:
        doc.add_heading("Proposed Level of Effort", level=2)
        doc.add_paragraph(loe)

    experience = str(toc.get("technical_experience_summary") or "").strip()
    if experience:
        doc.add_heading("Technical Experience and Past Performance References", level=2)
        doc.add_paragraph(experience)

    sample_outputs = str(toc.get("sample_outputs_summary") or "").strip()
    if sample_outputs:
        doc.add_heading("Sample Technical Outputs", level=2)
        doc.add_paragraph(sample_outputs)

    financial_summary = str(toc.get("financial_summary") or "").strip()
    if financial_summary:
        doc.add_heading("Financial Proposal Companion", level=2)
        doc.add_paragraph(financial_summary)

    cost_structure = toc.get("cost_structure")
    if isinstance(cost_structure, list) and cost_structure:
        doc.add_heading("Indicative Cost Structure", level=2)
        for row in cost_structure:
            if not isinstance(row, dict):
                continue
            cost_bucket = str(row.get("cost_bucket") or "Cost Bucket").strip()
            basis = str(row.get("basis") or "").strip()
            estimate = str(row.get("estimate") or "").strip()
            notes = str(row.get("notes") or "").strip()
            doc.add_paragraph(cost_bucket, style="List Bullet")
            details = []
            if basis:
                details.append(f"Basis: {basis}")
            if estimate:
                details.append(f"Estimate: {estimate}")
            if notes:
                details.append(f"Notes: {notes}")
            if details:
                doc.add_paragraph(" | ".join(details))

    pricing_assumptions = toc.get("pricing_assumptions")
    if isinstance(pricing_assumptions, list) and pricing_assumptions:
        doc.add_heading("Pricing Assumptions", level=2)
        for item in pricing_assumptions:
            doc.add_paragraph(str(item), style="List Bullet")

    payment_schedule = str(toc.get("payment_schedule_summary") or "").strip()
    if payment_schedule:
        doc.add_heading("Payment Schedule Summary", level=2)
        doc.add_paragraph(payment_schedule)

    submission_package = toc.get("submission_package_checklist")
    if isinstance(submission_package, list) and submission_package:
        doc.add_heading("Submission Package Completeness", level=2)
        for row in submission_package:
            if not isinstance(row, dict):
                continue
            artifact = str(row.get("artifact") or "Submission Artifact").strip()
            owner = str(row.get("owner") or "").strip()
            status = str(row.get("status") or "").strip()
            notes = str(row.get("notes") or "").strip()
            doc.add_paragraph(artifact, style="List Bullet")
            details = []
            if owner:
                details.append(f"Owner: {owner}")
            if status:
                details.append(f"Status: {status}")
            if notes:
                details.append(f"Notes: {notes}")
            if details:
                doc.add_paragraph(" | ".join(details))

    attachment_manifest = toc.get("attachment_manifest")
    if isinstance(attachment_manifest, list) and attachment_manifest:
        doc.add_heading("Attachment Manifest", level=2)
        for row in attachment_manifest:
            if not isinstance(row, dict):
                continue
            attachment = str(row.get("attachment") or "Attachment").strip()
            required_for = str(row.get("required_for") or "").strip()
            owner = str(row.get("owner") or "").strip()
            status = str(row.get("status") or "").strip()
            notes = str(row.get("notes") or "").strip()
            doc.add_paragraph(attachment, style="List Bullet")
            details = []
            if required_for:
                details.append(f"Required for: {required_for}")
            if owner:
                details.append(f"Owner: {owner}")
            if status:
                details.append(f"Status: {status}")
            if notes:
                details.append(f"Notes: {notes}")
            if details:
                doc.add_paragraph(" | ".join(details))

    deliverables = toc.get("deliverables")
    if isinstance(deliverables, list) and deliverables:
        doc.add_heading("Workplan & Deliverables", level=2)
        for row in deliverables:
            if not isinstance(row, dict):
                continue
            deliverable = str(row.get("deliverable") or "Deliverable").strip()
            timing = str(row.get("timing") or "").strip()
            purpose_text = str(row.get("purpose") or "").strip()
            title = deliverable if not timing else f"{deliverable} ({timing})"
            doc.add_paragraph(title, style="List Bullet")
            if purpose_text:
                doc.add_paragraph(purpose_text)
        _add_indicator_logframe_block(
            doc,
            indicators=output_focus or outcome_focus,
            label="Suggested delivery indicator rows",
        )

    annex_readiness = toc.get("annex_readiness")
    if isinstance(annex_readiness, list) and annex_readiness:
        doc.add_heading("Annex Readiness", level=2)
        for item in annex_readiness:
            doc.add_paragraph(str(item), style="List Bullet")

    compliance_matrix = toc.get("compliance_matrix")
    if isinstance(compliance_matrix, list) and compliance_matrix:
        doc.add_heading("Procurement Compliance Matrix", level=2)
        for row in compliance_matrix:
            if not isinstance(row, dict):
                continue
            requirement = str(row.get("requirement") or "Requirement").strip()
            response_section = str(row.get("response_section") or "").strip()
            evidence = str(row.get("evidence") or "").strip()
            status = str(row.get("status") or "").strip()
            notes = str(row.get("notes") or "").strip()
            doc.add_paragraph(requirement, style="List Bullet")
            details = []
            if response_section:
                details.append(f"Response section: {response_section}")
            if evidence:
                details.append(f"Evidence: {evidence}")
            if status:
                details.append(f"Status: {status}")
            if notes:
                details.append(f"Notes: {notes}")
            if details:
                doc.add_paragraph(" | ".join(details))

    assumptions_risks = toc.get("assumptions_risks")
    if isinstance(assumptions_risks, list) and assumptions_risks:
        doc.add_heading("Assumptions & Risks", level=2)
        for item in assumptions_risks:
            doc.add_paragraph(str(item), style="List Bullet")


def _add_mel_indicator_summary_section(doc: Document, logframe_draft: Optional[Dict[str, Any]]) -> None:
    if not isinstance(logframe_draft, dict):
        return
    indicators = logframe_draft.get("indicators")
    if not isinstance(indicators, list) or not indicators:
        return

    doc.add_heading("MEL Indicator Summary", level=1)
    for indicator in indicators:
        if not isinstance(indicator, dict):
            continue
        indicator_id = str(indicator.get("indicator_id") or "").strip()
        name = str(indicator.get("name") or "").strip() or "Untitled indicator"
        title = f"{indicator_id} — {name}" if indicator_id else name
        doc.add_paragraph(title, style="List Bullet")

        result_level = str(indicator.get("result_level") or "").strip()
        baseline = str(indicator.get("baseline") or "").strip()
        target = str(indicator.get("target") or "").strip()
        frequency = str(indicator.get("frequency") or "").strip()
        formula = str(indicator.get("formula") or "").strip()
        definition = str(indicator.get("definition") or "").strip()
        data_source = str(indicator.get("data_source") or "").strip()
        means_of_verification = str(indicator.get("means_of_verification") or "").strip()
        owner = str(indicator.get("owner") or "").strip()
        citation = str(indicator.get("citation") or "").strip()
        justification = str(indicator.get("justification") or "").strip()
        disaggregation_raw = indicator.get("disaggregation")
        disaggregation = ""
        if isinstance(disaggregation_raw, list):
            disaggregation = ", ".join(str(item).strip() for item in disaggregation_raw if str(item).strip())
        elif isinstance(disaggregation_raw, str):
            disaggregation = disaggregation_raw.strip()

        details: list[str] = []
        if result_level:
            details.append(f"Result level: {result_level}")
        if baseline or target:
            details.append(f"Baseline/Target: {baseline or '-'} -> {target or '-'}")
        if frequency:
            details.append(f"Frequency: {frequency}")
        if formula:
            details.append(f"Formula: {formula}")
        if definition:
            details.append(f"Definition: {definition}")
            details.append(f"Result focus: {definition[:140].rstrip()}{'...' if len(definition) > 140 else ''}")
        if data_source:
            details.append(f"Data source: {data_source}")
        if means_of_verification:
            details.append(f"Means of verification: {means_of_verification}")
        if owner:
            details.append(f"Owner: {owner}")
        if disaggregation:
            details.append(f"Disaggregation: {disaggregation}")
        if citation:
            details.append(f"Citation: {citation}")
        if justification:
            details.append(f"Justification: {justification}")
            details.append(
                f"Measurement intent: {justification[:140].rstrip()}{'...' if len(justification) > 140 else ''}"
            )
        if details:
            doc.add_paragraph(" | ".join(details))


def _add_template_profile_section(doc: Document, profile: Dict[str, Any]) -> None:
    doc.add_heading("Template Profile", level=1)
    doc.add_paragraph(f"Template: {profile.get('template_display_name')} ({profile.get('template_key')})")
    doc.add_paragraph(f"Coverage rate: {float(profile.get('coverage_rate') or 0.0):.0%}")

    raw_required = profile.get("required_sections")
    required = list(raw_required) if isinstance(raw_required, list) else []
    raw_present = profile.get("present_sections")
    present = list(raw_present) if isinstance(raw_present, list) else []
    raw_missing = profile.get("missing_sections")
    missing = list(raw_missing) if isinstance(raw_missing, list) else []
    if required:
        doc.add_paragraph(f"Required sections: {', '.join(str(x) for x in required)}")
        doc.add_paragraph(f"Present sections: {', '.join(str(x) for x in present) or '-'}")
        if missing:
            doc.add_paragraph(f"Missing sections: {', '.join(str(x) for x in missing)}")


def _add_export_contract_section(doc: Document, contract: Dict[str, Any]) -> None:
    doc.add_heading("Export Contract Check", level=1)
    status = str(contract.get("status") or "warning").upper()
    doc.add_paragraph(f"Status: {status}")

    missing_sections = contract.get("missing_required_sections")
    if isinstance(missing_sections, list) and missing_sections:
        doc.add_paragraph(f"Missing required ToC sections: {', '.join(str(x) for x in missing_sections)}")
    else:
        doc.add_paragraph("Missing required ToC sections: none")

    expected_docx_headings = contract.get("expected_docx_headings")
    if isinstance(expected_docx_headings, list) and expected_docx_headings:
        doc.add_paragraph(f"Expected donor headings: {', '.join(str(x) for x in expected_docx_headings)}")

    missing_headers = contract.get("missing_required_primary_sheet_headers")
    if isinstance(missing_headers, list) and missing_headers:
        doc.add_paragraph(f"Missing required primary sheet headers: {', '.join(str(x) for x in missing_headers)}")

    warnings = contract.get("warnings")
    if isinstance(warnings, list) and warnings:
        doc.add_paragraph(f"Warnings: {', '.join(str(x) for x in warnings)}")


def build_docx_from_toc(
    toc_draft: Dict[str, Any],
    donor_id: str,
    logframe_draft: Optional[Dict[str, Any]] = None,
    citations: Optional[List[Dict[str, Any]]] = None,
    critic_findings: Optional[List[Dict[str, Any]]] = None,
    review_comments: Optional[List[Dict[str, Any]]] = None,
    quality_summary: Optional[Dict[str, Any]] = None,
) -> bytes:
    """Конвертирует ToC draft в форматированный .docx."""
    doc = Document()

    title = doc.add_heading(f"Theory of Change — {donor_id}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Donor: {donor_id}")
    doc.add_paragraph("Generated by: GrantFlow")

    toc_content = _toc_root(toc_draft, donor_id)
    profile = build_export_template_profile(donor_id=donor_id, toc_payload=toc_content)
    contract = evaluate_export_contract(donor_id=donor_id, toc_payload=toc_content)
    _add_template_profile_section(doc, profile)
    _add_export_contract_section(doc, contract)
    _add_quality_summary_section(doc, quality_summary or {})
    _add_review_readiness_section(
        doc,
        quality_summary=quality_summary or {},
        citations=citations or toc_draft.get("citations") or [],
        critic_findings=critic_findings or [],
        review_comments=review_comments or [],
    )

    donor_key = resolve_export_template_key(donor_id=donor_id, toc_payload=toc_draft)
    if donor_key == "evaluation_rfq":
        _render_evaluation_rfq_toc(doc, toc_content, logframe_draft=logframe_draft)
    elif donor_key == "usaid":
        _render_usaid_toc(doc, toc_content, logframe_draft=logframe_draft)
    elif donor_key == "eu":
        _render_eu_toc(doc, toc_content, logframe_draft=logframe_draft)
    elif donor_key == "worldbank":
        _render_worldbank_toc(doc, toc_content, logframe_draft=logframe_draft)
    elif donor_key == "giz":
        _render_giz_toc(doc, toc_content, logframe_draft=logframe_draft)
    elif donor_key == "un_agencies":
        _render_un_agencies_toc(doc, toc_content, logframe_draft=logframe_draft)
    elif donor_key in {"state_department", "us_state_department", "u.s. department of state", "us department of state"}:
        _render_state_department_toc(doc, toc_content, logframe_draft=logframe_draft)
    else:
        _render_generic_toc(doc, toc_content)

    _add_mel_indicator_summary_section(doc, logframe_draft)
    _add_citation_traceability_section(doc, citations or toc_draft.get("citations") or [])
    _add_critic_findings_section(doc, critic_findings or [])
    _add_review_comments_section(doc, review_comments or [])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def save_docx_to_file(
    toc_draft: Dict[str, Any],
    donor_id: str,
    output_path: str,
    logframe_draft: Optional[Dict[str, Any]] = None,
    citations: Optional[List[Dict[str, Any]]] = None,
    critic_findings: Optional[List[Dict[str, Any]]] = None,
    review_comments: Optional[List[Dict[str, Any]]] = None,
    quality_summary: Optional[Dict[str, Any]] = None,
) -> str:
    """Сохраняет .docx на диск."""
    content = build_docx_from_toc(
        toc_draft,
        donor_id,
        logframe_draft=logframe_draft,
        citations=citations,
        critic_findings=critic_findings,
        review_comments=review_comments,
        quality_summary=quality_summary,
    )
    with open(output_path, "wb") as f:
        f.write(content)
    return output_path
