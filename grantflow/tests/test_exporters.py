import json
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from grantflow.exporters.donor_contracts import evaluate_export_contract
from grantflow.exporters.excel_builder import build_xlsx_from_logframe
from grantflow.exporters.word_builder import build_docx_from_toc


FIXTURES_DIR = Path(__file__).parent / "fixtures"


def _fixture_json(name: str):
    path = FIXTURES_DIR / name
    return json.loads(path.read_text(encoding="utf-8"))


def _sample_citations():
    return [
        {
            "stage": "mel",
            "citation_type": "rag_result",
            "namespace": "usaid_ads201",
            "source": "/tmp/usaid_guide.pdf",
            "page": 12,
            "chunk": 3,
            "chunk_id": "usaid_ads201_p12_c0",
            "used_for": "EG.3.2-1",
            "label": "USAID ADS 201 p.12",
            "excerpt": "Official indicator guidance excerpt",
            "citation_confidence": 0.82,
            "result_level": "outcome",
            "statement_path": "toc.development_objectives[0].description",
            "evidence_signal": "PMP and indicator reference evidence",
            "review_hint": "Use this citation to support the indicator row in the USAID monitoring package.",
        },
        {
            "stage": "architect",
            "citation_type": "strategy_reference",
            "namespace": "usaid_ads201",
            "source": "strategy::usaid",
            "page": "",
            "chunk": "",
            "chunk_id": "strategy::usaid::toc",
            "used_for": "toc_claim",
            "label": "Strategy reference",
            "excerpt": "",
            "citation_confidence": 0.75,
            "result_level": "general",
            "statement_path": "toc.project_goal",
            "evidence_signal": "strategy guidance only",
            "review_hint": "Use as strategy context only; replace with grounded evidence before external review.",
        },
    ]


def _sample_critic_findings():
    return [
        {
            "finding_id": "finding-123",
            "status": "acknowledged",
            "severity": "high",
            "section": "toc",
            "code": "TOC_SCHEMA_INVALID",
            "review_title": "ToC Schema Invalid",
            "review_bucket": "compliance",
            "triage_priority": "high",
            "message": "ToC schema contract mismatch.",
            "recommended_action": "Revise the ToC so it matches the donor review package before the next export.",
            "reviewer_next_step": "ToC Schema Invalid: Revise the ToC so it matches the donor review package before the next export.",
            "fix_hint": "Revise objective structure.",
            "version_id": "toc_v2",
            "source": "rules",
        }
    ]


def _sample_review_comments():
    return [
        {
            "comment_id": "comment-123",
            "status": "resolved",
            "section": "toc",
            "author": "reviewer-1",
            "message": "Adjusted objective wording and assumptions.",
            "version_id": "toc_v2",
            "linked_finding_id": "finding-123",
            "ts": "2026-02-25T10:00:00Z",
        }
    ]


def _sample_quality_summary():
    return {
        "quality_score": 8.5,
        "critic_score": 8.0,
        "needs_revision": False,
        "engine": "rules+llm",
        "rule_score": 8.0,
        "llm_score": 8.5,
        "fatal_flaw_count": 1,
        "citation_count": 4,
    }


def test_word_export_includes_citation_traceability_section():
    toc_draft = {
        "toc": {
            "brief": "Sample ToC brief",
            "objectives": [{"title": "Obj 1", "description": "Desc", "citation": "usaid_ads201"}],
        }
    }
    content = build_docx_from_toc(
        toc_draft,
        "usaid",
        citations=_sample_citations(),
        critic_findings=_sample_critic_findings(),
        review_comments=_sample_review_comments(),
        quality_summary=_sample_quality_summary(),
    )
    doc = Document(BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Citation Traceability" in text
    assert "USAID ADS 201 p.12" in text
    assert "Official indicator guidance excerpt" in text
    assert "conf 0.82" in text
    assert "level outcome" in text
    assert "path toc.development_objectives[0].description" in text
    assert "evidence PMP and indicator reference evidence" in text
    assert "Review hint: Use this citation to support the indicator row in the USAID monitoring package." in text
    assert "Critic Findings" in text
    assert "ToC Schema Invalid" in text
    assert "Review Comments" in text
    assert "Adjusted objective wording and assumptions." in text
    assert "Quality Summary" in text
    assert "Quality score: 8.5" in text
    assert "Fatal flaw count: 1" in text
    assert "Review Readiness" in text
    assert "Open critic findings: 1" in text
    assert "Acknowledged critic findings: 1" in text
    assert "Resolved critic findings: 0" in text
    assert "High-severity open findings: 1" in text
    assert "Open review comments: 0" in text
    assert "Fallback/strategy citations: 1" in text
    assert "Acknowledged review comments: 0" in text
    assert "Resolved review comments: 1" in text
    assert "Critic finding resolution rate:" in text
    assert "Critic finding acknowledgment rate:" in text
    assert "Reviewer workflow resolution rate:" in text
    assert "Reviewer workflow acknowledgment rate:" in text
    assert "Next primary review action: resolve_finding" in text
    assert "Finding ack queue: 0" in text
    assert "Finding resolve queue: 1" in text
    assert "Comment ack queue: 0" in text
    assert "Comment resolve queue: 0" in text
    assert "Comment reopen queue: 1" in text
    assert "Comment threads aged 3-7d:" in text
    assert "Top reviewer action 1:" in text
    assert "Reviewer next step:" in text
    assert "Recommended action:" in text


def test_word_export_uses_donor_specific_sections_for_usaid_eu_worldbank():
    usaid_toc = {
        "toc": {
            "project_goal": "Improve civic services",
            "development_objectives": [
                {
                    "do_id": "DO1",
                    "description": "Improved digital delivery",
                    "intermediate_results": [
                        {
                            "ir_id": "IR1.1",
                            "description": "Civil service capacity strengthened",
                            "outputs": [
                                {
                                    "output_id": "O1.1.1",
                                    "description": "Training delivered",
                                    "indicators": [
                                        {
                                            "indicator_code": "EG.1-1",
                                            "name": "Officials trained",
                                            "target": "300",
                                            "justification": "Capacity milestone",
                                            "citation": "USAID ADS 201 p.12",
                                        }
                                    ],
                                }
                            ],
                        }
                    ],
                }
            ],
            "critical_assumptions": ["Government commitment remains stable"],
        }
    }
    eu_toc = {
        "toc": {
            "overall_objective": {"objective_id": "OO1", "title": "Digital governance", "rationale": "EU fit"},
            "specific_objectives": [
                {"objective_id": "SO1", "title": "Service quality", "rationale": "Improve delivery standards"}
            ],
            "expected_outcomes": [
                {"outcome_id": "OUT1", "title": "Citizen trust", "expected_change": "Higher service satisfaction"}
            ],
            "assumptions": ["Policy continuity"],
            "risks": ["Election-year disruption"],
        }
    }
    wb_toc = {
        "toc": {
            "project_development_objective": "Improve municipal service delivery",
            "objectives": [
                {"objective_id": "PDO1", "title": "Service delivery quality", "description": "Better outcomes"}
            ],
            "results_chain": [
                {
                    "result_id": "R1",
                    "title": "Digital workflows adopted",
                    "description": "Local agencies use digital case-management",
                    "indicator_focus": "Processing time",
                }
            ],
            "assumptions": ["Budget continuity"],
            "risks": ["Procurement delays"],
        }
    }

    donor_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_001",
                "name": "Institutional service adoption score",
                "result_level": "outcome",
                "baseline": "0 institutions",
                "target": "4 institutions",
                "frequency": "quarterly",
                "formula": "Count of institutions meeting service adoption criteria",
                "definition": "Tracks institutional service adoption as an outcome-level implementation result.",
                "justification": "Maps service adoption outcome into a monitored proposal review row.",
                "means_of_verification": "Verification annexes and partner review notes",
                "owner": "Project M&E manager",
            },
            {
                "indicator_id": "IND_002",
                "name": "Civil service modernization learning uptake",
                "result_level": "impact",
                "baseline": "0 ministries",
                "target": "2 ministries",
                "frequency": "annual",
                "definition": "Tracks modernization learning uptake as an impact-level governance result.",
                "justification": "Maps learning uptake into a goal-level review and verification row.",
                "means_of_verification": "CLA evidence and review memos",
                "owner": "MEL lead",
            },
        ]
    }

    usaid_doc = Document(BytesIO(build_docx_from_toc(usaid_toc, "usaid", logframe_draft=donor_logframe)))
    eu_doc = Document(BytesIO(build_docx_from_toc(eu_toc, "eu", logframe_draft=donor_logframe)))
    state_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_003",
                "name": "Independent media resilience score",
                "result_level": "impact",
                "means_of_verification": "Editorial risk logs and resilience reviews",
                "owner": "Program manager and media partner leads",
            },
            {
                "indicator_id": "IND_004",
                "name": "Agency performance monitoring score",
                "result_level": "outcome",
                "means_of_verification": "ISR aide-memoires and results framework updates",
                "owner": "PIU results lead",
            },
        ]
    }

    wb_doc = Document(BytesIO(build_docx_from_toc(wb_toc, "worldbank", logframe_draft=state_logframe)))
    usaid_text = "\n".join(p.text for p in usaid_doc.paragraphs)
    eu_text = "\n".join(p.text for p in eu_doc.paragraphs)
    wb_text = "\n".join(p.text for p in wb_doc.paragraphs)

    assert "USAID Results Framework" in usaid_text
    assert "Critical Assumptions" in usaid_text
    assert "Suggested performance monitoring focus:" in usaid_text
    assert "Suggested performance indicator rows:" in usaid_text
    assert "Baseline/Target: 0 institutions -> 4 institutions" in usaid_text
    assert "Result focus:" in usaid_text
    assert "Measurement intent:" in usaid_text
    assert "CLA evidence and review memos" in usaid_text
    assert "EU Intervention Logic" in eu_text
    assert "Overall Objective" in eu_text
    assert "Specific Objectives" in eu_text
    assert "Expected Outcomes" in eu_text
    assert "Suggested monitoring focus:" in eu_text
    assert "Suggested objective-level LogFrame rows:" in eu_text
    assert "Frequency: quarterly" in eu_text
    assert "Verification annexes and partner review notes" in eu_text
    assert "Assumptions" in eu_text
    assert "Risks" in eu_text
    assert "World Bank Results Framework" in wb_text
    assert "Project Development Objective (PDO)" in wb_text
    assert "PDO1" in wb_text
    assert "Results Chain" in wb_text
    assert "Suggested PDO monitoring focus:" in wb_text
    assert "Suggested PDO indicator rows:" in wb_text
    assert "ISR aide-memoires and results framework updates" in wb_text


def test_word_export_uses_donor_specific_sections_for_giz_and_state_department():
    giz_toc = {
        "toc": {
            "programme_objective": "Strengthen SME resilience",
            "outputs": ["Coaching delivered", "Partner toolkits produced"],
            "outcomes": [
                {
                    "title": "SMEs improve recovery planning",
                    "description": "Participating SMEs institutionalize basic continuity plans",
                    "partner_role": "Chamber of Commerce delivers coaching",
                }
            ],
            "sustainability_factors": ["Partner co-financing"],
            "assumptions_risks": ["Macroeconomic volatility"],
        }
    }
    state_toc = {
        "toc": {
            "strategic_context": "Independent media under pressure",
            "program_goal": "Improve resilience of local media ecosystem",
            "objectives": [
                {
                    "objective": "Increase legal preparedness",
                    "line_of_effort": "rights",
                    "expected_change": "Media organizations handle legal pressure better",
                }
            ],
            "stakeholder_map": ["Local media outlets", "Legal aid groups"],
            "risk_mitigation": ["Anonymous reporting channels"],
        }
    }

    giz_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_GIZ_1",
                "name": "SME continuity adoption score",
                "result_level": "outcome",
                "baseline": "0 SMEs",
                "target": "40 SMEs",
                "frequency": "semiannual",
                "formula": "Count of SMEs meeting continuity adoption criteria",
                "means_of_verification": "Partner coaching records and SME verification visits",
                "owner": "Programme M&E lead and chamber focal points",
            },
            {
                "indicator_id": "IND_GIZ_2",
                "name": "Coaching completion rate",
                "result_level": "output",
                "baseline": "0 firms",
                "target": "60 firms",
                "frequency": "quarterly",
                "formula": "Count of firms completing coaching package",
                "means_of_verification": "Attendance registers and delivery logs",
                "owner": "Technical assistance team lead",
            },
        ]
    }

    giz_doc = Document(BytesIO(build_docx_from_toc(giz_toc, "giz", logframe_draft=giz_logframe)))
    state_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_005",
                "name": "Media resilience score",
                "result_level": "impact",
                "means_of_verification": "Editorial risk logs and resilience reviews",
                "owner": "Program manager and media partner leads",
            }
        ]
    }

    state_doc = Document(BytesIO(build_docx_from_toc(state_toc, "state_department", logframe_draft=state_logframe)))
    giz_text = "\n".join(p.text for p in giz_doc.paragraphs)
    state_text = "\n".join(p.text for p in state_doc.paragraphs)

    assert "GIZ Results & Sustainability Logic" in giz_text
    assert "Programme Objective" in giz_text
    assert "Sustainability Factors" in giz_text
    assert "Assumptions & Risks" in giz_text
    assert "Suggested implementation monitoring focus:" in giz_text
    assert "Suggested delivery verification focus:" in giz_text
    assert "Suggested programme-objective indicator rows:" in giz_text
    assert "Suggested output indicator rows:" in giz_text
    assert "Suggested outcome indicator rows:" in giz_text
    assert "Baseline/Target: 0 SMEs -> 40 SMEs" in giz_text
    assert "Partner coaching records and SME verification visits" in giz_text
    assert "U.S. Department of State Program Logic" in state_text
    assert "Strategic Context" in state_text
    assert "Risk Mitigation" in state_text
    assert "Suggested strategic monitoring focus:" in state_text
    assert "Editorial risk logs and resilience reviews" in state_text


def test_word_export_uses_un_agencies_template_and_focus_bridge():
    un_toc = {
        "toc": {
            "brief": "Inclusive education access remains uneven across target districts.",
            "objectives": [
                {
                    "title": "Improve inclusive education system readiness",
                    "description": "District education actors adopt inclusive planning and referral practices.",
                    "citation": "UNICEF education note",
                }
            ],
            "outcomes": [
                {
                    "title": "More children with disabilities access adapted services",
                    "description": "Referral and classroom support pathways operate consistently.",
                }
            ],
        }
    }
    un_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_UN_1",
                "name": "Inclusive service readiness score",
                "result_level": "outcome",
                "baseline": "0 districts",
                "target": "12 districts",
                "frequency": "annual",
                "formula": "Count of districts meeting inclusive service readiness criteria",
                "means_of_verification": "Partner monitoring records and school support checklists",
                "owner": "Programme manager and inclusion focal points",
            }
        ]
    }

    un_doc = Document(BytesIO(build_docx_from_toc(un_toc, "un_agencies", logframe_draft=un_logframe)))
    un_text = "\n".join(p.text for p in un_doc.paragraphs)

    assert "UN Agency Program Logic" in un_text
    assert "Overview" in un_text
    assert "Development Objectives" in un_text
    assert "Suggested monitoring focus:" in un_text
    assert "Suggested overview indicator rows:" in un_text
    assert "Suggested objective indicator rows:" in un_text
    assert "Suggested outcome indicator rows:" in un_text
    assert "Baseline/Target: 0 districts -> 12 districts" in un_text
    assert "Partner monitoring records and school support checklists" in un_text


def test_word_export_includes_template_profile_and_missing_sections_summary():
    eu_toc_incomplete = {
        "toc": {
            "overall_objective": {"objective_id": "OO1", "title": "Digital governance", "rationale": "EU fit"},
        }
    }
    doc = Document(BytesIO(build_docx_from_toc(eu_toc_incomplete, "eu")))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Template Profile" in text
    assert "EU Intervention Logic (eu)" in text
    assert "Coverage rate: 33%" in text
    assert "Missing sections: specific_objectives, expected_outcomes" in text


def test_word_export_includes_mel_indicator_summary_when_logframe_provided():
    toc_draft = {"toc": {"brief": "Sample ToC brief"}}
    logframe_draft = {
        "indicators": [
            {
                "indicator_id": "IND_001",
                "name": "Service coverage rate",
                "result_level": "outcome",
                "baseline": "0%",
                "target": "30%",
                "frequency": "quarterly",
                "formula": "(Numerator / Denominator) * 100",
                "definition": "Share of target population receiving the service.",
                "data_source": "PMP indicator tracking dataset",
                "means_of_verification": "Verified PMP records and spot-check files",
                "owner": "MEL lead and implementing partner M&E team",
                "disaggregation": ["sex", "age", "location"],
                "citation": "USAID ADS 201 p.12",
                "justification": "Tracks outcome-level adoption.",
            }
        ]
    }

    doc = Document(BytesIO(build_docx_from_toc(toc_draft, "usaid", logframe_draft=logframe_draft)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "MEL Indicator Summary" in text
    assert "IND_001 — Service coverage rate" in text
    assert "Result level: outcome" in text
    assert "Baseline/Target: 0% -> 30%" in text
    assert "Frequency: quarterly" in text
    assert "Data source: PMP indicator tracking dataset" in text
    assert "Means of verification: Verified PMP records and spot-check files" in text
    assert "Owner: MEL lead and implementing partner M&E team" in text
    assert "Disaggregation: sex, age, location" in text


def test_word_export_includes_export_contract_section():
    eu_toc_incomplete = {
        "toc": {
            "overall_objective": {"objective_id": "OO1", "title": "Digital governance", "rationale": "EU fit"},
        }
    }
    doc = Document(BytesIO(build_docx_from_toc(eu_toc_incomplete, "eu")))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Export Contract Check" in text
    assert "Status: WARNING" in text
    assert "Missing required ToC sections: specific_objectives, expected_outcomes" in text
    assert (
        "Expected donor headings: EU Intervention Logic, Overall Objective, Specific Objectives, Expected Outcomes"
        in text
    )


def test_excel_export_includes_citations_sheet():
    logframe_draft = {
        "indicators": [
            {
                "indicator_id": "IND_001",
                "name": "Indicator A",
                "justification": "Test",
                "citation": "USAID ADS 201 p.12",
                "baseline": "0",
                "target": "100",
            }
        ]
    }
    content = build_xlsx_from_logframe(
        logframe_draft,
        "usaid",
        citations=_sample_citations(),
        critic_findings=_sample_critic_findings(),
        review_comments=_sample_review_comments(),
        quality_summary=_sample_quality_summary(),
    )
    wb = load_workbook(BytesIO(content))
    assert "LogFrame" in wb.sheetnames
    assert "Quality Summary" in wb.sheetnames
    assert "Review Readiness" in wb.sheetnames
    assert "Citations" in wb.sheetnames
    assert "Critic Findings" in wb.sheetnames
    assert "Review Comments" in wb.sheetnames

    ws = wb["Citations"]
    rows = list(ws.iter_rows(values_only=True))
    assert rows[0][:7] == ("Stage", "Type", "Used For", "Label", "Confidence", "Result Level", "Statement Path")
    headers = [str(cell or "") for cell in rows[0]]
    assert "Evidence Signal" in headers
    assert "Review Hint" in headers
    assert any(row[3] == "USAID ADS 201 p.12" for row in rows[1:])
    assert any(abs(float(row[4]) - 0.82) < 1e-9 for row in rows[1:] if row[4] is not None)
    assert any(row[5] == "outcome" for row in rows[1:])
    assert any(row[6] == "toc.development_objectives[0].description" for row in rows[1:])
    evidence_idx = headers.index("Evidence Signal")
    review_hint_idx = headers.index("Review Hint")
    assert any(row[evidence_idx] == "PMP and indicator reference evidence" for row in rows[1:])
    assert any(
        row[review_hint_idx] == "Use this citation to support the indicator row in the USAID monitoring package."
        for row in rows[1:]
    )

    findings_rows = list(wb["Critic Findings"].iter_rows(values_only=True))
    assert findings_rows[0][:4] == ("Status", "Severity", "Section", "Code")
    assert any(row[3] == "TOC_SCHEMA_INVALID" for row in findings_rows[1:])
    findings_headers = [str(cell or "") for cell in findings_rows[0]]
    assert "Review Title" in findings_headers
    assert "Recommended Action" in findings_headers
    assert "Reviewer Next Step" in findings_headers

    comments_rows = list(wb["Review Comments"].iter_rows(values_only=True))
    assert comments_rows[0][:4] == ("Status", "Section", "Author", "Message")
    assert any(row[6] == "2026-02-25T10:00:00Z" for row in comments_rows[1:])

    quality_rows = list(wb["Quality Summary"].iter_rows(values_only=True))
    quality_map = {str(row[0]): row[1] for row in quality_rows[1:] if row and row[0]}
    assert quality_map["Quality score"] == 8.5
    assert quality_map["Critic engine"] == "rules+llm"
    assert quality_map["Fatal flaw count"] == 1

    readiness_rows = list(wb["Review Readiness"].iter_rows(values_only=True))
    readiness_map = {str(row[0]): row[1] for row in readiness_rows[1:] if row and row[0]}
    assert readiness_map["Open critic findings"] == 1
    assert readiness_map["Acknowledged critic findings"] == 1
    assert readiness_map["Resolved critic findings"] == 0
    assert readiness_map["High-severity open findings"] == 1
    assert readiness_map["Open review comments"] == 0
    assert readiness_map["Acknowledged review comments"] == 0
    assert readiness_map["Resolved review comments"] == 1
    assert readiness_map["Fallback/strategy citations"] == 1
    assert readiness_map["Next primary review action"] == "resolve_finding"
    assert readiness_map["Finding ack queue"] == 0
    assert readiness_map["Finding resolve queue"] == 1
    assert readiness_map["Comment ack queue"] == 0
    assert readiness_map["Comment resolve queue"] == 0
    assert readiness_map["Comment reopen queue"] == 1
    assert "Critic finding resolution rate" in readiness_map
    assert "Critic finding acknowledgment rate" in readiness_map
    assert "Reviewer workflow resolution rate" in readiness_map
    assert "Reviewer workflow acknowledgment rate" in readiness_map
    assert "Comment threads aged 3-7d" in readiness_map
    assert "Top reviewer action 1" in readiness_map


def test_excel_export_logframe_sheet_includes_smart_indicator_columns():
    logframe_draft = {
        "indicators": [
            {
                "indicator_id": "IND_001",
                "name": "Service coverage rate",
                "result_level": "outcome",
                "toc_statement_path": "toc.outcomes[0]",
                "justification": "Tracks outcome-level service adoption.",
                "citation": "USAID ADS 201 p.12",
                "baseline": "0%",
                "target": "30%",
                "frequency": "quarterly",
                "formula": "(Numerator / Denominator) * 100",
                "definition": "Share of target beneficiaries receiving digital service.",
                "data_source": "PMP indicator tracking dataset",
                "disaggregation": ["sex", "age", "location"],
                "means_of_verification": "Verified PMP records and spot-check files",
                "owner": "MEL lead and implementing partner M&E team",
                "evidence_excerpt": "Service uptake evidence extracted from validated implementation records.",
            }
        ]
    }
    wb = load_workbook(BytesIO(build_xlsx_from_logframe(logframe_draft, "usaid")))
    ws = wb["LogFrame"]
    rows = list(ws.iter_rows(values_only=True))

    assert rows[0] == (
        "Indicator ID",
        "Name",
        "Result Level",
        "ToC Statement Path",
        "Justification",
        "Citation",
        "Readiness Hint",
        "Result Focus",
        "Measurement Intent",
        "Baseline",
        "Target",
        "Frequency",
        "Formula",
        "Definition",
        "Data Source",
        "Disaggregation",
        "Means of Verification",
        "Owner",
        "Evidence Excerpt",
    )
    assert rows[1][0] == "IND_001"
    assert rows[1][2] == "outcome"
    assert rows[1][3] == "toc.outcomes[0]"
    assert rows[1][6] == "review-ready"
    assert "Share of target beneficiaries receiving digital service." in str(rows[1][7])
    assert "Tracks outcome-level service adoption." in str(rows[1][8])
    assert rows[1][11] == "quarterly"
    assert rows[1][12] == "(Numerator / Denominator) * 100"
    assert rows[1][14] == "PMP indicator tracking dataset"
    assert rows[1][15] == "sex, age, location"
    assert rows[1][16] == "Verified PMP records and spot-check files"
    assert rows[1][17] == "MEL lead and implementing partner M&E team"
    assert "validated implementation records" in str(rows[1][18])


def test_excel_export_includes_template_meta_sheet():
    eu_toc_incomplete = {
        "toc": {
            "overall_objective": {"objective_id": "OO1", "title": "Digital governance", "rationale": "EU fit"},
        }
    }
    content = build_xlsx_from_logframe({"indicators": []}, "eu", toc_draft=eu_toc_incomplete)
    wb = load_workbook(BytesIO(content))
    assert "Template Meta" in wb.sheetnames

    rows = list(wb["Template Meta"].iter_rows(values_only=True))
    row_map = {str(row[0]): row[1] for row in rows[1:] if row and row[0]}
    assert row_map["Template Key"] == "eu"
    assert row_map["Template Display"] == "EU Intervention Logic"
    assert float(row_map["Coverage Rate"]) < 0.34
    assert float(row_map["Coverage Rate"]) > 0.32
    assert "specific_objectives" in str(row_map["Missing Sections"])
    assert "expected_outcomes" in str(row_map["Missing Sections"])


def test_excel_donor_sheets_include_suggested_monitoring_focus_columns():
    donor_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_001",
                "name": "Institutional service adoption score",
                "result_level": "outcome",
                "means_of_verification": "Verification annexes and partner review notes",
                "owner": "Project M&E manager",
            },
            {
                "indicator_id": "IND_002",
                "name": "Civil service modernization learning uptake",
                "result_level": "impact",
                "means_of_verification": "CLA evidence and review memos",
                "owner": "MEL lead",
            },
        ]
    }
    usaid_toc = {
        "toc": {
            "project_goal": "Improve civic services",
            "development_objectives": [
                {
                    "do_id": "DO1",
                    "description": "Improved digital delivery",
                    "intermediate_results": [
                        {
                            "ir_id": "IR1.1",
                            "description": "Capacity strengthened",
                            "outputs": [
                                {
                                    "output_id": "O1.1.1",
                                    "description": "Training delivered",
                                    "indicators": [],
                                }
                            ],
                        }
                    ],
                }
            ],
        }
    }
    eu_toc = {
        "toc": {
            "overall_objective": {"objective_id": "OO1", "title": "Digital governance", "rationale": "EU fit"},
            "specific_objectives": [
                {"objective_id": "SO1", "title": "Service quality", "rationale": "Improve delivery standards"}
            ],
            "expected_outcomes": [
                {"outcome_id": "OUT1", "title": "Citizen trust", "expected_change": "Higher service satisfaction"}
            ],
        }
    }

    usaid_wb = load_workbook(BytesIO(build_xlsx_from_logframe(donor_logframe, "usaid", toc_draft=usaid_toc)))
    eu_wb = load_workbook(BytesIO(build_xlsx_from_logframe(donor_logframe, "eu", toc_draft=eu_toc)))

    usaid_rows = list(usaid_wb["USAID_RF"].iter_rows(values_only=True))
    eu_rows = list(eu_wb["EU_Intervention"].iter_rows(values_only=True))

    assert "Suggested Monitoring Focus" in usaid_rows[0]
    assert "Suggested Means of Verification" in usaid_rows[0]
    assert "Suggested Owner" in usaid_rows[0]
    assert "Suggested Result Focus" in usaid_rows[0]
    assert "Suggested Measurement Intent" in usaid_rows[0]
    assert "Institutional service adoption score" in str(usaid_rows[1])
    assert "Project M&E manager" in str(usaid_rows[1])

    assert "Suggested Monitoring Focus" in eu_rows[0]
    assert "Suggested Means of Verification" in eu_rows[0]
    assert "Suggested Owner" in eu_rows[0]
    assert "Suggested Result Focus" in eu_rows[0]
    assert "Suggested Measurement Intent" in eu_rows[0]
    assert "Institutional service adoption score" in str(eu_rows[1])
    assert "Verification annexes and partner review notes" in str(eu_rows[1])


def test_excel_export_includes_export_contract_sheet():
    usaid_toc = {
        "toc": {
            "project_goal": "Improve civic services",
            "development_objectives": [
                {
                    "do_id": "DO1",
                    "description": "Improved digital delivery",
                    "intermediate_results": [],
                }
            ],
        }
    }
    content = build_xlsx_from_logframe({"indicators": []}, "usaid", toc_draft=usaid_toc)
    wb = load_workbook(BytesIO(content))
    assert "Export Contract" in wb.sheetnames
    rows = list(wb["Export Contract"].iter_rows(values_only=True))
    row_map = {str(row[0]): row[1] for row in rows[1:] if row and row[0]}
    assert row_map["Status"] == "pass"
    assert row_map["Missing Required ToC Sections"] == "-"
    assert "LogFrame" in str(row_map["Required Workbook Sheets"])
    assert "USAID_RF" in str(row_map["Required Workbook Sheets"])
    assert row_map["Expected Primary Sheet"] == "USAID_RF"
    assert "DO ID" in str(row_map["Expected Primary Sheet Headers"])
    assert "IR Description" in str(row_map["Actual Primary Sheet Headers"])
    assert row_map["Missing Required Primary Sheet Headers"] == "-"


def test_export_contract_docx_mode_skips_workbook_sheet_requirements():
    contract = evaluate_export_contract(
        donor_id="usaid",
        toc_payload={
            "project_goal": "Improve civic services",
            "development_objectives": [{"do_id": "DO1", "description": "Improved digital delivery"}],
        },
    )
    assert contract["workbook_validation_enabled"] is False
    assert contract["status"] == "pass"
    assert contract["missing_required_sheets"] == []


def test_export_contract_detects_missing_primary_sheet_headers():
    contract = evaluate_export_contract(
        donor_id="usaid",
        toc_payload={
            "project_goal": "Improve civic services",
            "development_objectives": [{"do_id": "DO1", "description": "Improved digital delivery"}],
        },
        workbook_sheetnames=["LogFrame", "USAID_RF", "Template Meta"],
        workbook_primary_sheet_headers=["DO ID", "DO Description", "IR ID"],
    )
    assert contract["status"] == "warning"
    assert contract["missing_required_sheets"] == []
    assert "IR Description" in contract["missing_required_primary_sheet_headers"]
    assert "missing_required_primary_sheet_headers" in contract["warnings"]


def test_export_contract_matches_golden_snapshot():
    fixture = _fixture_json("export_contract_golden.json")
    cases = fixture.get("cases") if isinstance(fixture, dict) else None
    assert isinstance(cases, list) and cases
    for case in cases:
        assert isinstance(case, dict)
        payload = case.get("input")
        expected = case.get("expected")
        assert isinstance(payload, dict)
        assert isinstance(expected, dict)
        contract = evaluate_export_contract(**payload)
        for key, value in expected.items():
            assert contract.get(key) == value


def test_exporters_accept_critic_finding_id_alias():
    findings = [
        {
            "id": "finding-alias-1",
            "status": "open",
            "severity": "medium",
            "section": "general",
            "code": "FINDING_ALIAS",
            "message": "Alias id should be rendered in exports.",
            "fix_suggestion": "Use canonical finding id field.",
            "source": "rules",
        }
    ]

    docx_content = build_docx_from_toc({"toc": {"brief": "Alias test"}}, "usaid", critic_findings=findings)
    doc = Document(BytesIO(docx_content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "finding_id=finding-alias-1" in text

    xlsx_content = build_xlsx_from_logframe({"indicators": []}, "usaid", critic_findings=findings)
    wb = load_workbook(BytesIO(xlsx_content))
    rows = list(wb["Critic Findings"].iter_rows(values_only=True))
    headers = [str(cell or "") for cell in rows[0]]
    finding_id_index = headers.index("Finding ID")
    assert any(row[finding_id_index] == "finding-alias-1" for row in rows[1:])


def test_excel_export_includes_donor_specific_sheets():
    usaid_toc = {
        "toc": {
            "development_objectives": [
                {
                    "do_id": "DO1",
                    "description": "Improved digital delivery",
                    "intermediate_results": [
                        {
                            "ir_id": "IR1.1",
                            "description": "Civil service capacity strengthened",
                            "outputs": [
                                {
                                    "output_id": "O1.1.1",
                                    "description": "Training delivered",
                                    "indicators": [
                                        {
                                            "indicator_code": "EG.1-1",
                                            "name": "Officials trained",
                                            "target": "300",
                                            "justification": "Capacity milestone",
                                            "citation": "USAID ADS 201 p.12",
                                        }
                                    ],
                                }
                            ],
                        }
                    ],
                }
            ]
        }
    }
    eu_toc = {
        "toc": {
            "overall_objective": {"objective_id": "OO1", "title": "Digital governance", "rationale": "EU fit"},
            "specific_objectives": [
                {"objective_id": "SO1", "title": "Service quality", "rationale": "Improve delivery standards"}
            ],
            "expected_outcomes": [
                {"outcome_id": "OUT1", "title": "Citizen trust", "expected_change": "Higher service satisfaction"}
            ],
            "assumptions": ["Policy continuity"],
            "risks": ["Election-year disruption"],
        }
    }
    wb_toc = {
        "toc": {
            "project_development_objective": "Improve municipal service delivery",
            "objectives": [
                {"objective_id": "PDO1", "title": "Service delivery quality", "description": "Better outcomes"}
            ],
            "results_chain": [
                {
                    "result_id": "R1",
                    "title": "Digital workflows adopted",
                    "description": "Local agencies use digital case-management",
                    "indicator_focus": "Processing time",
                }
            ],
        }
    }
    giz_toc = {
        "toc": {
            "programme_objective": "Strengthen SME resilience",
            "outputs": ["Coaching delivered"],
            "outcomes": [
                {
                    "title": "SMEs improve recovery planning",
                    "description": "Participating SMEs institutionalize continuity plans",
                    "partner_role": "Chamber of Commerce",
                }
            ],
            "sustainability_factors": ["Partner co-financing"],
            "assumptions_risks": ["Macroeconomic volatility"],
        }
    }
    state_toc = {
        "toc": {
            "strategic_context": "Independent media under pressure",
            "program_goal": "Improve media ecosystem resilience",
            "objectives": [
                {
                    "objective": "Increase legal preparedness",
                    "line_of_effort": "rights",
                    "expected_change": "Media organizations handle pressure better",
                }
            ],
            "stakeholder_map": ["Local media outlets"],
            "risk_mitigation": ["Anonymous reporting channels"],
        }
    }

    usaid_wb = load_workbook(BytesIO(build_xlsx_from_logframe({"indicators": []}, "usaid", toc_draft=usaid_toc)))
    eu_wb = load_workbook(BytesIO(build_xlsx_from_logframe({"indicators": []}, "eu", toc_draft=eu_toc)))
    donor_focus_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_WB_1",
                "name": "Agency performance monitoring score",
                "result_level": "outcome",
                "baseline": "0 agencies",
                "target": "8 agencies",
                "frequency": "semi-annual",
                "formula": "Count of agencies meeting performance score threshold",
                "means_of_verification": "ISR aide-memoires and results framework updates",
                "owner": "PIU results lead",
            },
            {
                "indicator_id": "IND_STATE_1",
                "name": "Media resilience score",
                "result_level": "impact",
                "baseline": "0 organizations",
                "target": "6 organizations",
                "frequency": "annual",
                "means_of_verification": "Editorial risk logs and resilience reviews",
                "owner": "Program manager and media partner leads",
            },
        ]
    }

    wb_wb = load_workbook(BytesIO(build_xlsx_from_logframe(donor_focus_logframe, "worldbank", toc_draft=wb_toc)))
    giz_focus_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_GIZ_1",
                "name": "SME continuity adoption score",
                "result_level": "outcome",
                "baseline": "0 SMEs",
                "target": "40 SMEs",
                "frequency": "semiannual",
                "formula": "Count of SMEs meeting continuity adoption criteria",
                "means_of_verification": "Partner coaching records and SME verification visits",
                "owner": "Programme M&E lead and chamber focal points",
            },
            {
                "indicator_id": "IND_GIZ_2",
                "name": "Coaching completion rate",
                "result_level": "output",
                "baseline": "0 firms",
                "target": "60 firms",
                "frequency": "quarterly",
                "formula": "Count of firms completing coaching package",
                "means_of_verification": "Attendance registers and delivery logs",
                "owner": "Technical assistance team lead",
            },
        ]
    }
    un_toc = {
        "toc": {
            "brief": "Inclusive education access remains uneven across target districts.",
            "objectives": [
                {
                    "title": "Improve inclusive education system readiness",
                    "description": "District education actors adopt inclusive planning and referral practices.",
                }
            ],
            "outcomes": [
                {
                    "title": "More children with disabilities access adapted services",
                    "description": "Referral and classroom support pathways operate consistently.",
                }
            ],
        }
    }
    un_focus_logframe = {
        "indicators": [
            {
                "indicator_id": "IND_UN_1",
                "name": "Inclusive service readiness score",
                "result_level": "outcome",
                "baseline": "0 districts",
                "target": "12 districts",
                "frequency": "annual",
                "formula": "Count of districts meeting inclusive service readiness criteria",
                "means_of_verification": "Partner monitoring records and school support checklists",
                "owner": "Programme manager and inclusion focal points",
            }
        ]
    }
    giz_wb = load_workbook(BytesIO(build_xlsx_from_logframe(giz_focus_logframe, "giz", toc_draft=giz_toc)))
    un_wb = load_workbook(BytesIO(build_xlsx_from_logframe(un_focus_logframe, "un_agencies", toc_draft=un_toc)))
    state_wb = load_workbook(
        BytesIO(build_xlsx_from_logframe(donor_focus_logframe, "state_department", toc_draft=state_toc))
    )

    assert "USAID_RF" in usaid_wb.sheetnames
    assert "EU_Intervention" in eu_wb.sheetnames
    assert "EU_Assumptions_Risks" in eu_wb.sheetnames
    assert "WB_Results" in wb_wb.sheetnames
    assert "GIZ_Results" in giz_wb.sheetnames
    assert "UN_Results" in un_wb.sheetnames
    assert "StateDept_Results" in state_wb.sheetnames

    usaid_rows = list(usaid_wb["USAID_RF"].iter_rows(values_only=True))
    assert usaid_rows[0][0] == "DO ID"
    assert any(row[0] == "DO1" and row[6] == "EG.1-1" for row in usaid_rows[1:])

    eu_rows = list(eu_wb["EU_Intervention"].iter_rows(values_only=True))
    assert eu_rows[0][:2] == ("Level", "ID")
    assert any(row[0] == "Specific Objective" and row[1] == "SO1" for row in eu_rows[1:])
    assert any(row[0] == "Outcome" and row[1] == "OUT1" for row in eu_rows[1:])
    eu_aux_rows = list(eu_wb["EU_Assumptions_Risks"].iter_rows(values_only=True))
    assert any(row[0] == "Assumption" for row in eu_aux_rows[1:])
    assert any(row[0] == "Risk" for row in eu_aux_rows[1:])

    wb_rows = list(wb_wb["WB_Results"].iter_rows(values_only=True))
    assert wb_rows[0][0] == "Level"
    assert any(row[0] == "PDO" for row in wb_rows[1:])
    assert any(row[0] == "Result" and row[1] == "R1" for row in wb_rows[1:])
    assert "Suggested Monitoring Focus" in wb_rows[0]
    assert "Suggested Means of Verification" in wb_rows[0]
    assert "Suggested Owner" in wb_rows[0]
    assert "Suggested Baseline -> Target" in wb_rows[0]
    assert "Suggested Frequency" in wb_rows[0]
    assert "Suggested Formula" in wb_rows[0]
    assert any("ISR aide-memoires and results framework updates" in str(row) for row in wb_rows[1:])
    assert any("0 agencies -> 8 agencies" in str(row) for row in wb_rows[1:])

    giz_rows = list(giz_wb["GIZ_Results"].iter_rows(values_only=True))
    assert any(row[0] == "Programme Objective" for row in giz_rows[1:])
    assert any(row[0] == "Outcome" for row in giz_rows[1:])
    assert "Suggested Monitoring Focus" in giz_rows[0]
    assert "Suggested Means of Verification" in giz_rows[0]
    assert "Suggested Owner" in giz_rows[0]
    assert "Suggested Baseline -> Target" in giz_rows[0]
    assert "Suggested Frequency" in giz_rows[0]
    assert "Suggested Formula" in giz_rows[0]
    assert "Suggested Result Focus" in giz_rows[0]
    assert "Suggested Measurement Intent" in giz_rows[0]
    assert any("Attendance registers and delivery logs" in str(row) for row in giz_rows[1:])
    assert any("0 firms -> 60 firms" in str(row) for row in giz_rows[1:])

    un_rows = list(un_wb["UN_Results"].iter_rows(values_only=True))
    assert un_rows[0][:4] == ("Level", "Title", "Description", "Suggested Monitoring Focus")
    assert "Suggested Baseline -> Target" in un_rows[0]
    assert "Suggested Frequency" in un_rows[0]
    assert "Suggested Formula" in un_rows[0]
    assert "Suggested Result Focus" in un_rows[0]
    assert "Suggested Measurement Intent" in un_rows[0]
    assert any(row[0] == "Overview" for row in un_rows[1:])
    assert any(
        row[0] == "Objective" and row[1] == "Improve inclusive education system readiness" for row in un_rows[1:]
    )
    assert any("Partner monitoring records and school support checklists" in str(row) for row in un_rows[1:])
    assert any("0 districts -> 12 districts" in str(row) for row in un_rows[1:])

    state_rows = list(state_wb["StateDept_Results"].iter_rows(values_only=True))
    assert any(row[0] == "Strategic Context" for row in state_rows[1:])
    assert any(row[0] == "Objective" for row in state_rows[1:])
    assert "Suggested Monitoring Focus" in state_rows[0]
    assert "Suggested Means of Verification" in state_rows[0]
    assert "Suggested Owner" in state_rows[0]
    assert "Suggested Result Focus" in state_rows[0]
    assert "Suggested Measurement Intent" in state_rows[0]
    assert any("Editorial risk logs and resilience reviews" in str(row) for row in state_rows[1:])


def test_export_contract_normalizes_usaid_alias_payload_keys():
    contract = evaluate_export_contract(
        donor_id="usaid",
        toc_payload={
            "goal": "Improve service quality",
            "objectives": [
                {
                    "id": "OBJ-1",
                    "title": "Digital workflows adopted",
                    "results": [{"id": "R1", "title": "Civil servants trained"}],
                }
            ],
            "assumptions": ["Budget continuity"],
        },
    )
    assert contract["status"] == "pass"
    assert contract["missing_required_sections"] == []
    assert contract["template_key"] == "usaid"


def test_worldbank_exporters_accept_variant_payload_shapes():
    wb_toc = {
        "toc": {
            "pdo": "Increase municipal service performance",
            "development_objectives": [
                {"id": "OBJ1", "title": "Service quality improvement", "expected_change": "Faster response times"}
            ],
            "outcomes": [
                {"outcome_id": "OUT1", "name": "Backlog reduced", "description": "Average queue time declines"}
            ],
            "assumptions_risks": ["Procurement delays"],
        }
    }

    wb_doc = Document(BytesIO(build_docx_from_toc(wb_toc, "worldbank")))
    wb_text = "\n".join(p.text for p in wb_doc.paragraphs)
    assert "World Bank Results Framework" in wb_text
    assert "Project Development Objective (PDO)" in wb_text
    assert "OBJ1" in wb_text
    assert "OUT1" in wb_text

    wb_xlsx = load_workbook(BytesIO(build_xlsx_from_logframe({"indicators": []}, "worldbank", toc_draft=wb_toc)))
    rows = list(wb_xlsx["WB_Results"].iter_rows(values_only=True))
    assert any(row[0] == "PDO" and "Increase municipal service performance" in str(row[3]) for row in rows[1:])
    assert any(row[0] == "Objective" and row[1] == "OBJ1" for row in rows[1:])
    assert any(row[0] == "Result" and row[1] == "OUT1" for row in rows[1:])


def test_exporters_compact_noisy_architect_toc_text_in_word_and_excel():
    toc_payload = {
        "toc": {
            "project_development_objective": (
                "Improve public sector performance and service delivery in Uzbekistan. "
                "Evidence hint: present. Grounding gate warning: architect_retrieval_no_hits."
            ),
            "objectives": [
                {
                    "objective_id": "OBJ1",
                    "title": "Strengthen institutional performance for public sector performance and service delivery",
                    "description": (
                        "Improve implementation reliability for public sector performance and service delivery priorities. "
                        "Evidence hint: present. Replace repeated boilerplate."
                    ),
                }
            ],
            "results_chain": [
                {
                    "result_id": "R1",
                    "title": "Agencies adopt operational improvements for public sector performance and service delivery",
                    "description": (
                        "Participating agencies implement workflow improvements that increase the quality of "
                        "public sector performance and service delivery delivery."
                    ),
                    "indicator_focus": "Processing time and institutional compliance",
                }
            ],
        }
    }

    wb_doc = Document(BytesIO(build_docx_from_toc(toc_payload, "worldbank")))
    wb_text = "\n".join(p.text for p in wb_doc.paragraphs)
    assert "Evidence hint:" not in wb_text
    assert "architect_retrieval_no_hits" not in wb_text
    assert "service delivery delivery" not in wb_text
    assert "Improve public sector performance and service delivery in Uzbekistan" in wb_text

    wb_xlsx = load_workbook(BytesIO(build_xlsx_from_logframe({"indicators": []}, "worldbank", toc_draft=toc_payload)))
    rows = list(wb_xlsx["WB_Results"].iter_rows(values_only=True))
    rendered = "\n".join(" | ".join("" if cell is None else str(cell) for cell in row) for row in rows[1:])
    assert "Evidence hint:" not in rendered
    assert "architect_retrieval_no_hits" not in rendered
    assert "service delivery delivery" not in rendered
    assert "Improve public sector performance and service delivery in Uzbekistan" in rendered


def test_exporters_support_evaluation_rfq_mode():
    toc_draft = {
        "proposal_mode": "evaluation_rfq",
        "toc": {
            "proposal_mode": "evaluation_rfq",
            "rfq_profile": "katch_final_assessment",
            "brief": "Technical response for an external project performance evaluation.",
            "background_context": "Assignment covers project performance and learning needs.",
            "evaluation_purpose": "Assess outcome-level change and implementation performance.",
            "organization_information": "Registered organization with audited financials and Central Asia operating status.",
            "technical_approach_summary": "Mixed-methods design, document review, sampling, analysis software, risks, and ethics.",
            "sampling_plan": "Stakeholder and beneficiary sampling across project geographies.",
            "analytical_software": ["Qualitative coding software"],
            "ethical_considerations": ["Informed consent", "Confidentiality", "Do no harm"],
            "evaluation_questions": [
                "What changes can be substantiated?",
                "What recommendations should guide follow-on action?",
            ],
            "methodology_overview": "Mixed-methods evaluation design with triangulated evidence review.",
            "methodology_components": [
                {
                    "method": "Outcome Harvesting",
                    "purpose": "Capture substantiated outcome-level change",
                    "respondent_group": "Stakeholders and partners",
                    "evidence_source": "Validated project records",
                }
            ],
            "team_composition": [
                {"role": "Team Lead", "responsibility": "Lead technical quality and client communication"}
            ],
            "key_personnel": [
                {
                    "name": "Proposed Team Lead",
                    "role": "Team Lead",
                    "qualifications": "Senior evaluation lead with donor reporting experience",
                    "level_of_effort": "Lead oversight across inception and reporting",
                    "cv_status": "ready",
                }
            ],
            "level_of_effort_summary": "Activity-based person-days by phase and role.",
            "technical_experience_summary": "Comparable evaluation experience in development-sector final assessments.",
            "sample_outputs_summary": "Annexed final evaluation reports from comparable assignments.",
            "compliance_matrix": [
                {
                    "requirement": "Organization information and legal status package",
                    "response_section": "Organization Information",
                    "evidence": "Registration certificate and audited financials",
                    "status": "ready",
                    "notes": "Attached in annex package",
                }
            ],
            "deliverables": [
                {
                    "deliverable": "Inception Report",
                    "timing": "Week 1",
                    "purpose": "Confirm evaluation matrix and work plan",
                }
            ],
            "workplan_summary": ["Mobilize, validate design, and execute fieldwork."],
            "assumptions_risks": ["Stakeholder access remains available."],
            "annex_readiness": ["CVs", "Registration certificate", "Sample outputs", "LOE matrix"],
        },
    }
    logframe_draft = {
        "proposal_mode": "evaluation_rfq",
        "indicators": [
            {
                "indicator_id": "IND_001",
                "name": "Deliverable milestone: Inception Report",
                "result_level": "output",
                "baseline": "0 deliverables",
                "target": "1 deliverable",
                "frequency": "bi-weekly",
                "formula": "Count of required deliverables completed and accepted against the agreed evaluation work plan",
                "definition": "Tracks whether the inception package is completed with reviewer-ready documentation.",
                "justification": "Maps the deliverable schedule into an evaluation RFQ management indicator.",
                "means_of_verification": "Accepted deliverable package and QA checklist",
                "owner": "Evaluation team lead and operations coordinator",
            }
        ],
    }

    doc = Document(BytesIO(build_docx_from_toc(toc_draft, "un_agencies", logframe_draft=logframe_draft)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Evaluation RFQ Technical Proposal" in text
    assert "Organization Information" in text
    assert "Assignment Background" in text
    assert "Evaluation Purpose" in text
    assert "Analysis and Proposed Approaches / Methodologies" in text
    assert "Sampling Plan" in text
    assert "Personnel and Team Composition" in text
    assert "Key Personnel and CV Readiness" in text
    assert "Senior evaluation lead with donor reporting experience" in text
    assert "CV status: ready" in text
    assert "Proposed Level of Effort" in text
    assert "Technical Experience and Past Performance References" in text
    assert "Procurement Compliance Matrix" in text
    assert "Registration certificate and audited financials" in text
    assert "Workplan & Deliverables" in text

    wb = load_workbook(BytesIO(build_xlsx_from_logframe(logframe_draft, "un_agencies", toc_draft=toc_draft)))
    assert "Evaluation_Plan" in wb.sheetnames
    headers = [cell.value for cell in wb["Evaluation_Plan"][1]]
    assert headers == ["Section", "ID", "Title", "Description"]
    rendered = "\n".join(
        " | ".join("" if cell is None else str(cell) for cell in row)
        for row in wb["Evaluation_Plan"].iter_rows(values_only=True)
    )
    assert "Organization Information" in rendered
    assert "Analysis and Proposed Approaches / Methodologies" in rendered
    assert "Proposed Team Lead (Team Lead)" in rendered
    assert "CV: ready" in rendered
    assert "Proposed Level of Effort" in rendered
    assert "Technical Experience and Past Performance References" in rendered
    assert "Organization information and legal status package" in rendered
    assert "Registration certificate and audited financials" in rendered
    contract = evaluate_export_contract(
        donor_id="un_agencies",
        toc_payload=toc_draft,
        workbook_sheetnames=wb.sheetnames,
        workbook_primary_sheet_headers=headers,
    )
    assert contract["template_key"] == "evaluation_rfq"
    assert contract["status"] == "pass"
